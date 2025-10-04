"""
测试高级Markdown提取功能
"""
import os
import tempfile
from docx import Document

def create_complex_test_document():
    """创建一个复杂的测试Word文档"""
    doc = Document()
    
    # 添加各种格式的内容
    doc.add_heading('复杂测试文档', 0)
    doc.add_heading('第一章 介绍', 1)
    doc.add_heading('1.1 背景信息', 2)
    
    # 添加段落
    p1 = doc.add_paragraph('这是一个普通段落，包含一些文本。')
    p2 = doc.add_paragraph('这是包含')
    p2.add_run('粗体文本').bold = True
    p2.add_run('和')
    p2.add_run('斜体文本').italic = True
    p2.add_run('的段落。')
    
    # 添加列表
    doc.add_paragraph('主要特点：', style='List Bullet')
    doc.add_paragraph('特点一：高质量翻译', style='List Bullet')
    doc.add_paragraph('特点二：格式保持', style='List Bullet')
    doc.add_paragraph('特点三：智能识别', style='List Bullet')
    
    # 添加编号列表
    doc.add_paragraph('实施步骤：', style='List Number')
    doc.add_paragraph('第一步：上传文档', style='List Number')
    doc.add_paragraph('第二步：选择语言', style='List Number')
    doc.add_paragraph('第三步：下载结果', style='List Number')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '功能'
    hdr_cells[1].text = '描述'
    hdr_cells[2].text = '状态'
    
    # 数据行
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '翻译'
    row1_cells[1].text = '多语言翻译'
    row1_cells[2].text = '完成'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = '格式保持'
    row2_cells[1].text = '保持原格式'
    row2_cells[2].text = '完成'
    
    # 保存到临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name

def test_advanced_extraction():
    """测试高级提取功能"""
    print("测试高级Markdown提取功能...")
    
    # 创建复杂测试文档
    test_file = create_complex_test_document()
    
    try:
        # 导入我们的函数
        import sys
        sys.path.append('.')
        from app import word_to_markdown_advanced, word_to_markdown_direct, word_to_markdown_simple
        
        print("\n=== 测试高级提取方法 ===")
        advanced_result = word_to_markdown_advanced(test_file)
        if advanced_result:
            print("高级提取成功")
            print(f"内容长度: {len(advanced_result)} 字符")
            print("提取结果预览:")
            print("-" * 50)
            print(advanced_result[:500] + "..." if len(advanced_result) > 500 else advanced_result)
            print("-" * 50)
        else:
            print("高级提取失败")
        
        print("\n=== 测试直接解析方法 ===")
        direct_result = word_to_markdown_direct(test_file)
        if direct_result:
            print("直接解析成功")
            print(f"内容长度: {len(direct_result)} 字符")
        else:
            print("直接解析失败")
        
        print("\n=== 测试简化解析方法 ===")
        simple_result = word_to_markdown_simple(test_file)
        if simple_result:
            print("简化解析成功")
            print(f"内容长度: {len(simple_result)} 字符")
        else:
            print("简化解析失败")
        
        # 比较结果
        print("\n=== 结果比较 ===")
        results = [
            ("高级提取", advanced_result),
            ("直接解析", direct_result),
            ("简化解析", simple_result)
        ]
        
        for name, result in results:
            if result:
                print(f"{name}: {len(result)} 字符")
            else:
                print(f"{name}: 失败")
        
    except Exception as e:
        print(f"测试失败: {e}")
    finally:
        # 清理临时文件
        try:
            os.unlink(test_file)
        except:
            pass

if __name__ == "__main__":
    test_advanced_extraction()
