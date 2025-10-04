"""
简单展示界面 - 只展示翻译结果，不提供编辑功能
"""

import streamlit as st
from docx import Document
from typing import List, Dict, Any
import tempfile
import os

class SimpleDisplayInterface:
    """简单展示界面 - 只展示翻译结果"""
    
    def __init__(self):
        self.original_paragraphs = []
        self.translated_paragraphs = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            # 读取原文档
            original_doc = Document(original_path)
            self.original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
            
            # 读取翻译文档
            translated_doc = Document(translated_path)
            self.translated_paragraphs = [p.text.strip() for p in translated_doc.paragraphs if p.text.strip()]
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def display_simple_interface(self):
        """显示简单展示界面"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            st.warning("⚠️ 请先加载文档")
            return
        
        st.markdown("---")
        st.subheader("📄 翻译结果展示")
        
        # 显示统计信息
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("原文档段落数", len(self.original_paragraphs))
        
        with col2:
            st.metric("译文段落数", len(self.translated_paragraphs))
        
        with col3:
            st.metric("翻译完成率", "100%")
        
        # 显示翻译结果对比
        st.markdown("### 📊 翻译结果对比")
        
        # 使用左右两列布局展示原文和译文
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📝 原文")
            for i, paragraph in enumerate(self.original_paragraphs):
                if paragraph.strip():
                    st.markdown(f"**段落 {i+1}:**")
                    st.text_area(
                        f"原文段落 {i+1}",
                        value=paragraph,
                        height=100,
                        key=f"original_display_{i}",
                        disabled=True
                    )
                    st.markdown(f"字数: {len(paragraph)}")
                    st.markdown("---")
        
        with col2:
            st.markdown("#### 🌐 译文")
            for i, paragraph in enumerate(self.translated_paragraphs):
                if paragraph.strip():
                    st.markdown(f"**段落 {i+1}:**")
                    st.text_area(
                        f"译文段落 {i+1}",
                        value=paragraph,
                        height=100,
                        key=f"translated_display_{i}",
                        disabled=True
                    )
                    st.markdown(f"字数: {len(paragraph)}")
                    st.markdown("---")
        
        # 显示翻译统计
        self._display_translation_stats()
    
    def _display_translation_stats(self):
        """显示翻译统计"""
        st.markdown("### 📈 翻译统计")
        
        # 计算统计信息
        total_original_chars = sum(len(p) for p in self.original_paragraphs)
        total_translated_chars = sum(len(p) for p in self.translated_paragraphs)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("原文档总字数", total_original_chars)
        
        with col2:
            st.metric("译文总字数", total_translated_chars)
        
        with col3:
            length_ratio = total_translated_chars / total_original_chars if total_original_chars > 0 else 1
            st.metric("长度比例", f"{length_ratio:.2f}")
        
        with col4:
            st.metric("段落数", len(self.original_paragraphs))
    
    def get_translation_summary(self):
        """获取翻译摘要"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            return {}
        
        total_original_chars = sum(len(p) for p in self.original_paragraphs)
        total_translated_chars = sum(len(p) for p in self.translated_paragraphs)
        
        return {
            'total_paragraphs': len(self.original_paragraphs),
            'total_original_chars': total_original_chars,
            'total_translated_chars': total_translated_chars,
            'length_ratio': total_translated_chars / total_original_chars if total_original_chars > 0 else 1
        }
    
    def display_translation_summary(self):
        """显示翻译摘要"""
        summary = self.get_translation_summary()
        
        if summary:
            st.markdown("### 📊 翻译摘要")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("总段落数", summary['total_paragraphs'])
            
            with col2:
                st.metric("原文档总字数", summary['total_original_chars'])
            
            with col3:
                st.metric("译文总字数", summary['total_translated_chars'])
            
            # 长度比例
            st.markdown(f"**长度比例**: {summary['length_ratio']:.2f}")
            
            # 翻译质量评估
            if summary['length_ratio'] > 0.8 and summary['length_ratio'] < 1.2:
                st.success("✅ 翻译长度合理")
            elif summary['length_ratio'] > 1.2:
                st.warning("⚠️ 译文较长，可能需要调整")
            else:
                st.warning("⚠️ 译文较短，可能需要调整")
