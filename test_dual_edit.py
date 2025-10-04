"""
测试左右编辑界面
"""

import streamlit as st
from dual_edit_interface import DualEditInterface
from docx import Document
import tempfile
import os

def create_test_documents():
    """创建测试文档"""
    # 创建原文文档
    original_doc = Document()
    original_doc.add_heading('Test Document', 0)
    original_doc.add_paragraph('This is the first paragraph in English.')
    original_doc.add_paragraph('This is the second paragraph with more content.')
    original_doc.add_heading('Section 1', level=1)
    original_doc.add_paragraph('This is a section paragraph.')
    original_doc.add_paragraph('This is another paragraph with detailed information.')
    original_doc.add_paragraph('This is the final paragraph of the document.')
    
    # 创建译文文档
    translated_doc = Document()
    translated_doc.add_heading('测试文档', 0)
    translated_doc.add_paragraph('这是第一段中文内容。')
    translated_doc.add_paragraph('这是第二段中文内容，包含更多信息。')
    translated_doc.add_heading('章节 1', level=1)
    translated_doc.add_paragraph('这是一个章节段落。')
    translated_doc.add_paragraph('这是另一个包含详细信息的段落。')
    translated_doc.add_paragraph('这是文档的最后一段。')
    
    # 保存到临时文件
    original_path = tempfile.mktemp(suffix='.docx')
    translated_path = tempfile.mktemp(suffix='.docx')
    
    original_doc.save(original_path)
    translated_doc.save(translated_path)
    
    return original_path, translated_path

def main():
    """主测试函数"""
    st.set_page_config(
        page_title="左右编辑界面测试",
        page_icon="✏️",
        layout="wide"
    )
    
    st.title("✏️ 左右编辑界面测试")
    st.markdown("---")
    
    # 创建测试文档
    if st.button("🚀 创建测试文档并测试左右编辑界面"):
        with st.spinner("正在创建测试文档..."):
            original_path, translated_path = create_test_documents()
            
            # 初始化左右编辑界面
            edit_interface = DualEditInterface()
            
            # 加载文档
            if edit_interface.load_documents(original_path, translated_path):
                st.success("✅ 测试文档创建成功！")
                
                # 显示编辑摘要
                edit_interface.display_edit_summary()
                
                # 显示左右编辑界面
                edit_interface.display_dual_edit_interface()
                
                # 显示所有段落编辑
                edit_interface.display_all_paragraphs_edit()
                
                # 最终输出测试
                st.markdown("---")
                st.subheader("📤 最终输出测试")
                
                if st.button("📄 生成最终文档", type="primary"):
                    with st.spinner("正在生成最终文档..."):
                        final_output_path = tempfile.mktemp(suffix='.docx')
                        
                        if edit_interface.create_final_document(final_output_path):
                            st.success("✅ 最终文档生成成功！")
                            
                            # 读取最终文档
                            with open(final_output_path, 'rb') as f:
                                final_data = f.read()
                            
                            # 提供下载
                            st.download_button(
                                label="📥 下载最终文档",
                                data=final_data,
                                file_name="test_final_document.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            # 清理临时文件
                            try:
                                os.unlink(final_output_path)
                            except:
                                pass
                        else:
                            st.error("❌ 最终文档生成失败")
                
                # 清理临时文件
                try:
                    os.unlink(original_path)
                    os.unlink(translated_path)
                except:
                    pass
            else:
                st.error("❌ 测试文档创建失败")
    
    # 功能说明
    st.markdown("---")
    st.subheader("🎯 左右编辑界面功能")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### ✅ 核心功能")
        st.markdown("- 左右分开显示：原文在左，译文在右")
        st.markdown("- 可编辑译文：右侧译文可以编辑修改")
        st.markdown("- 实时统计：字数、字符数、比例等")
        st.markdown("- 编辑操作：保存、重置、复制等")
    
    with col2:
        st.markdown("### 📊 编辑功能")
        st.markdown("- 段落选择器：选择要编辑的段落")
        st.markdown("- 所有段落编辑：批量编辑所有段落")
        st.markdown("- 最终输出：生成修改后的文档")
        st.markdown("- 编辑统计：修改率、修改段落数等")
    
    # 测试说明
    st.markdown("---")
    st.subheader("🧪 测试步骤")
    st.markdown("1. 点击'创建测试文档并测试左右编辑界面'按钮")
    st.markdown("2. 查看编辑摘要统计信息")
    st.markdown("3. 使用段落选择器选择不同段落")
    st.markdown("4. 在右侧编辑框中修改译文内容")
    st.markdown("5. 使用编辑操作按钮（保存、重置、复制等）")
    st.markdown("6. 展开'所有段落编辑'进行批量编辑")
    st.markdown("7. 点击'生成最终文档'并下载")

if __name__ == "__main__":
    main()
