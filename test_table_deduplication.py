"""
测试表格去重功能
"""

import os
import tempfile
from docx import Document
from docx.shared import Pt
from smart_translator import StructuralParser, SemanticTranslator, SmartReconstructor

def create_test_table_document():
    """创建包含表格的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('表格去重测试文档', 0)
    
    # 添加段落
    doc.add_paragraph('这是一个测试段落，用于验证表格去重功能。')
    
    # 创建表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目名称'
    hdr_cells[1].text = '功能描述'
    hdr_cells[2].text = '状态'
    
    # 数据行
    row1 = table.rows[1].cells
    row1[0].text = '智能翻译'
    row1[1].text = '基于AI的文档翻译'
    row1[2].text = '已完成'
    
    row2 = table.rows[2].cells
    row2[0].text = '格式保持'
    row2[1].text = '保持原文档格式'
    row2[2].text = '开发中'
    
    # 保存文档
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name

def test_table_parsing():
    """测试表格解析去重"""
    print("测试表格解析去重...")
    
    # 创建测试文档
    test_doc_path = create_test_table_document()
    print(f"创建测试文档: {test_doc_path}")
    
    # 测试解析
    parser = StructuralParser()
    result = parser.parse_document(test_doc_path)
    
    if result:
        print("表格解析成功!")
        
        # 统计表格内容
        table_cells = [item for item in result['content_layer'] if item['type'] == 'table_cell']
        print(f"表格单元格数量: {len(table_cells)}")
        
        # 检查重复内容
        cell_texts = [cell['text'] for cell in table_cells if cell['text'].strip()]
        unique_texts = set(cell_texts)
        
        print(f"唯一文本数量: {len(unique_texts)}")
        print(f"重复内容数量: {len(cell_texts) - len(unique_texts)}")
        
        if len(cell_texts) - len(unique_texts) == 0:
            print("表格解析去重成功，无重复内容")
        else:
            print("表格解析仍有重复内容")
            print("重复的文本:")
            for text in cell_texts:
                if cell_texts.count(text) > 1:
                    print(f"  - {text}")
        
        # 显示表格内容
        print("\n表格内容:")
        for i, cell in enumerate(table_cells):
            if cell['text'].strip():
                print(f"  {i+1}. 表格{cell.get('table_index', 0)+1} 行{cell.get('row', 0)+1} 列{cell.get('col', 0)+1}: {cell['text']}")
    
    else:
        print("表格解析失败")
    
    # 清理
    try:
        os.unlink(test_doc_path)
    except:
        pass

def test_duplicate_detection():
    """测试重复内容检测"""
    print("\n测试重复内容检测...")
    
    # 模拟重复内容
    test_items = [
        {'text': '项目A', 'type': 'table_cell'},
        {'text': '项目B', 'type': 'table_cell'},
        {'text': '项目A', 'type': 'table_cell'},  # 重复
        {'text': '项目C', 'type': 'table_cell'},
        {'text': '项目B', 'type': 'table_cell'},  # 重复
    ]
    
    # 检测重复
    texts = [item['text'] for item in test_items]
    unique_texts = set(texts)
    duplicates = len(texts) - len(unique_texts)
    
    print(f"总项目数: {len(texts)}")
    print(f"唯一项目数: {len(unique_texts)}")
    print(f"重复项目数: {duplicates}")
    
    if duplicates > 0:
        print("发现重复内容:")
        for text in unique_texts:
            count = texts.count(text)
            if count > 1:
                print(f"  - '{text}' 重复 {count} 次")
    else:
        print("无重复内容")

def main():
    """主测试函数"""
    print("表格去重功能测试")
    print("=" * 50)
    
    # 测试表格解析去重
    test_table_parsing()
    
    # 测试重复内容检测
    test_duplicate_detection()
    
    print("\n" + "=" * 50)
    print("表格去重测试完成!")

if __name__ == "__main__":
    main()
