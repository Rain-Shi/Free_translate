"""
测试段落对比功能
"""

import streamlit as st
from simple_comparison import SimpleParagraphComparison
from docx import Document
import tempfile
import os

def create_test_documents():
    """创建测试文档"""
    # 创建原文文档
    original_doc = Document()
    original_doc.add_heading('Test Document', 0)
    original_doc.add_paragraph('This is the first paragraph in English.')
    original_doc.add_paragraph('This is the second paragraph with more content.')
    original_doc.add_heading('Section 1', level=1)
    original_doc.add_paragraph('This is a section paragraph.')
    original_doc.add_paragraph('This is another paragraph with detailed information.')
    original_doc.add_paragraph('This is the final paragraph of the document.')
    
    # 创建译文文档
    translated_doc = Document()
    translated_doc.add_heading('测试文档', 0)
    translated_doc.add_paragraph('这是第一段中文内容。')
    translated_doc.add_paragraph('这是第二段中文内容，包含更多信息。')
    translated_doc.add_heading('章节 1', level=1)
    translated_doc.add_paragraph('这是一个章节段落。')
    translated_doc.add_paragraph('这是另一个包含详细信息的段落。')
    translated_doc.add_paragraph('这是文档的最后一段。')
    
    # 保存到临时文件
    original_path = tempfile.mktemp(suffix='.docx')
    translated_path = tempfile.mktemp(suffix='.docx')
    
    original_doc.save(original_path)
    translated_doc.save(translated_path)
    
    return original_path, translated_path

def main():
    """主测试函数"""
    st.set_page_config(
        page_title="段落对比功能测试",
        page_icon="📖",
        layout="wide"
    )
    
    st.title("📖 段落对比功能测试")
    st.markdown("---")
    
    # 创建测试文档
    if st.button("🚀 创建测试文档并测试段落对比"):
        with st.spinner("正在创建测试文档..."):
            original_path, translated_path = create_test_documents()
            
            # 初始化段落对比器
            comparison = SimpleParagraphComparison()
            
            # 加载文档
            if comparison.load_documents(original_path, translated_path):
                st.success("✅ 测试文档创建成功！")
                
                # 显示文档摘要
                comparison.display_summary()
                
                # 显示段落对比
                comparison.display_comparison()
                
                # 显示所有段落概览
                comparison.display_all_paragraphs()
                
                # 清理临时文件
                try:
                    os.unlink(original_path)
                    os.unlink(translated_path)
                except:
                    pass
            else:
                st.error("❌ 测试文档创建失败")
    
    # 功能说明
    st.markdown("---")
    st.subheader("🎯 段落对比功能")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### ✅ 核心功能")
        st.markdown("- 段落选择器：选择要对比的段落")
        st.markdown("- 原文译文对比：并排显示原文和译文")
        st.markdown("- 统计信息：字数、字符数、比例等")
        st.markdown("- 所有段落概览：可展开查看所有段落")
    
    with col2:
        st.markdown("### 📊 统计功能")
        st.markdown("- 文档摘要：段落数、总字符数")
        st.markdown("- 段落统计：字数、字符数")
        st.markdown("- 对比分析：长度比例、词数比例")
        st.markdown("- 概览模式：快速浏览所有段落")
    
    # 测试说明
    st.markdown("---")
    st.subheader("🧪 测试步骤")
    st.markdown("1. 点击'创建测试文档并测试段落对比'按钮")
    st.markdown("2. 查看文档摘要统计信息")
    st.markdown("3. 使用段落选择器选择不同段落")
    st.markdown("4. 查看原文和译文的详细对比")
    st.markdown("5. 展开'所有段落概览'查看所有段落")

if __name__ == "__main__":
    main()
