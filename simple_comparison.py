"""
简洁段落对比功能
"""

import streamlit as st
from docx import Document
from typing import List, Dict, Any
import tempfile
import os

class SimpleParagraphComparison:
    """简洁段落对比器"""
    
    def __init__(self):
        self.original_paragraphs = []
        self.translated_paragraphs = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            # 读取原文档
            original_doc = Document(original_path)
            self.original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
            
            # 读取翻译文档
            translated_doc = Document(translated_path)
            self.translated_paragraphs = [p.text.strip() for p in translated_doc.paragraphs if p.text.strip()]
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def display_comparison(self):
        """显示段落对比"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            st.warning("⚠️ 请先加载文档")
            return
        
        st.markdown("---")
        st.subheader("📖 段落对比")
        
        # 显示统计信息
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("原文档段落数", len(self.original_paragraphs))
        
        with col2:
            st.metric("翻译文档段落数", len(self.translated_paragraphs))
        
        with col3:
            st.metric("对比段落数", min(len(self.original_paragraphs), len(self.translated_paragraphs)))
        
        # 段落选择器
        max_paragraphs = min(len(self.original_paragraphs), len(self.translated_paragraphs))
        
        if max_paragraphs > 0:
            selected_paragraph = st.selectbox(
                "选择要对比的段落",
                options=list(range(1, max_paragraphs + 1)),
                format_func=lambda x: f"段落 {x}",
                key="paragraph_selector"
            )
            
            if selected_paragraph:
                para_index = selected_paragraph - 1
                
                # 显示段落对比
                self._display_paragraph_comparison(para_index)
    
    def _display_paragraph_comparison(self, para_index: int):
        """显示单个段落对比"""
        st.markdown("---")
        st.markdown("### 🔍 段落详细对比")
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📝 原文**")
            original_text = self.original_paragraphs[para_index] if para_index < len(self.original_paragraphs) else ""
            st.text_area(
                "原文内容",
                value=original_text,
                height=200,
                key=f"original_text_{para_index}",
                disabled=True
            )
            
            # 原文统计
            st.markdown(f"**字数**: {len(original_text)}")
            st.markdown(f"**字符数**: {len(original_text.replace(' ', ''))}")
        
        with col2:
            st.markdown("**🌐 译文**")
            translated_text = self.translated_paragraphs[para_index] if para_index < len(self.translated_paragraphs) else ""
            st.text_area(
                "译文内容",
                value=translated_text,
                height=200,
                key=f"translated_text_{para_index}",
                disabled=True
            )
            
            # 译文统计
            st.markdown(f"**字数**: {len(translated_text)}")
            st.markdown(f"**字符数**: {len(translated_text.replace(' ', ''))}")
        
        # 对比统计
        if original_text and translated_text:
            self._display_comparison_stats(original_text, translated_text)
    
    def _display_comparison_stats(self, original_text: str, translated_text: str):
        """显示对比统计"""
        st.markdown("---")
        st.markdown("### 📊 对比统计")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            length_ratio = len(translated_text) / len(original_text) if original_text else 1
            st.metric("长度比例", f"{length_ratio:.2f}")
        
        with col2:
            word_count_orig = len(original_text.split())
            word_count_trans = len(translated_text.split())
            st.metric("原文词数", word_count_orig)
        
        with col3:
            st.metric("译文词数", word_count_trans)
        
        with col4:
            if word_count_orig > 0:
                word_ratio = word_count_trans / word_count_orig
                st.metric("词数比例", f"{word_ratio:.2f}")
    
    def display_all_paragraphs(self):
        """显示所有段落概览"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            return
        
        st.markdown("---")
        st.subheader("📋 所有段落概览")
        
        # 创建可滚动的段落列表
        max_paragraphs = min(len(self.original_paragraphs), len(self.translated_paragraphs))
        
        for i in range(max_paragraphs):
            with st.expander(f"段落 {i+1}", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**原文**")
                    st.text(self.original_paragraphs[i][:100] + "..." if len(self.original_paragraphs[i]) > 100 else self.original_paragraphs[i])
                
                with col2:
                    st.markdown("**译文**")
                    st.text(self.translated_paragraphs[i][:100] + "..." if len(self.translated_paragraphs[i]) > 100 else self.translated_paragraphs[i])
    
    def get_document_summary(self):
        """获取文档摘要"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            return {}
        
        return {
            'original_paragraphs': len(self.original_paragraphs),
            'translated_paragraphs': len(self.translated_paragraphs),
            'total_characters_original': sum(len(p) for p in self.original_paragraphs),
            'total_characters_translated': sum(len(p) for p in self.translated_paragraphs),
            'avg_length_original': sum(len(p) for p in self.original_paragraphs) / len(self.original_paragraphs) if self.original_paragraphs else 0,
            'avg_length_translated': sum(len(p) for p in self.translated_paragraphs) / len(self.translated_paragraphs) if self.translated_paragraphs else 0
        }
    
    def display_summary(self):
        """显示文档摘要"""
        summary = self.get_document_summary()
        
        if summary:
            st.markdown("### 📊 文档摘要")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("原文档段落数", summary['original_paragraphs'])
            
            with col2:
                st.metric("译文段落数", summary['translated_paragraphs'])
            
            with col3:
                st.metric("原文档总字符", summary['total_characters_original'])
            
            with col4:
                st.metric("译文总字符", summary['total_characters_translated'])
