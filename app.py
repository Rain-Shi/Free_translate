import streamlit as st
import os
import tempfile
from docx import Document
import markdown
import pypandoc
import openai
import io

# 确保pandoc可用
try:
    pypandoc.get_pandoc_version()
except OSError:
    st.error("Pandoc未找到，正在下载...")
    pypandoc.download_pandoc()

# 设置页面配置
st.set_page_config(
    page_title="Word文档翻译工具",
    page_icon="📄",
    layout="wide"
)

def extract_document_fonts(word_file_path):
    """提取Word文档的字体信息"""
    try:
        doc = Document(word_file_path)
        fonts_info = {
            'default_font': '宋体',  # 默认字体
            'default_size': 12,      # 默认字号
            'heading_fonts': {},     # 各级标题字体
            'paragraph_fonts': [],   # 段落字体
            'bold_fonts': [],        # 粗体字体
            'italic_fonts': []       # 斜体字体
        }
        
        # 分析文档中的字体
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                for run in paragraph.runs:
                    if run.font.name:
                        font_name = run.font.name
                        font_size = run.font.size.pt if run.font.size else 12
                        
                        # 记录字体信息
                        if paragraph.style.name.startswith('Heading'):
                            level = paragraph.style.name.split()[-1] if 'Heading' in paragraph.style.name else '1'
                            fonts_info['heading_fonts'][f'heading_{level}'] = {
                                'name': font_name,
                                'size': font_size
                            }
                        else:
                            fonts_info['paragraph_fonts'].append({
                                'name': font_name,
                                'size': font_size,
                                'bold': run.bold,
                                'italic': run.italic
                            })
                            
                            if run.bold:
                                fonts_info['bold_fonts'].append(font_name)
                            if run.italic:
                                fonts_info['italic_fonts'].append(font_name)
        
        # 确定主要字体
        if fonts_info['paragraph_fonts']:
            # 使用最常见的字体作为默认字体
            font_counts = {}
            for font_info in fonts_info['paragraph_fonts']:
                font_name = font_info['name']
                font_counts[font_name] = font_counts.get(font_name, 0) + 1
            
            if font_counts:
                fonts_info['default_font'] = max(font_counts, key=font_counts.get)
        
        return fonts_info
    except Exception as e:
        st.warning(f"字体提取失败: {str(e)}")
        return {
            'default_font': '宋体',
            'default_size': 12,
            'heading_fonts': {},
            'paragraph_fonts': [],
            'bold_fonts': [],
            'italic_fonts': []
        }

def word_to_markdown_direct(word_file_path):
    """使用python-docx直接解析Word文档为Markdown - 增强版"""
    try:
        doc = Document(word_file_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
                
            # 检查段落样式
            style_name = paragraph.style.name.lower()
            
            # 更精确的标题检测
            if any(keyword in style_name for keyword in ['heading 1', 'title', 'heading1']):
                markdown_content.append(f"# {text}")
            elif any(keyword in style_name for keyword in ['heading 2', 'heading2']):
                markdown_content.append(f"## {text}")
            elif any(keyword in style_name for keyword in ['heading 3', 'heading3']):
                markdown_content.append(f"### {text}")
            elif any(keyword in style_name for keyword in ['heading 4', 'heading4']):
                markdown_content.append(f"#### {text}")
            elif any(keyword in style_name for keyword in ['heading 5', 'heading5']):
                markdown_content.append(f"##### {text}")
            elif any(keyword in style_name for keyword in ['heading 6', 'heading6']):
                markdown_content.append(f"###### {text}")
            elif any(keyword in style_name for keyword in ['list', 'bullet', 'number']):
                # 检测列表类型
                if 'number' in style_name or 'numbered' in style_name:
                    markdown_content.append(f"1. {text}")
                else:
                    markdown_content.append(f"- {text}")
            else:
                # 高级格式检测和保持
                formatted_text = _format_paragraph_with_runs(paragraph)
                markdown_content.append(formatted_text)
        
        # 处理表格
        for table in doc.tables:
            markdown_content.append("")  # 空行分隔
            markdown_content.append(_table_to_markdown(table))
            markdown_content.append("")  # 空行分隔
        
        return "\n".join(markdown_content)
    except Exception as e:
        st.error(f"直接解析Word文档失败: {str(e)}")
        return None

def _format_paragraph_with_runs(paragraph):
    """格式化段落，保持粗体、斜体等格式"""
    if not paragraph.runs:
        return paragraph.text
    
    formatted_parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        # 应用格式
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        elif run.underline:
            text = f"<u>{text}</u>"
            
        formatted_parts.append(text)
    
    return ''.join(formatted_parts)

def _table_to_markdown(table):
    """将Word表格转换为Markdown表格 - 改进版"""
    if not table.rows:
        return ""
    
    markdown_table = []
    
    # 获取表格的最大列数，确保所有行都有相同的列数
    max_cols = max(len(row.cells) for row in table.rows) if table.rows else 0
    
    # 处理所有行，包括表头和数据行
    for i, row in enumerate(table.rows):
        cells = []
        for j in range(max_cols):
            if j < len(row.cells):
                # 获取单元格文本，保持格式
                cell_text = row.cells[j].text.strip()
                # 处理空单元格
                if not cell_text:
                    cell_text = " "
                cells.append(cell_text)
            else:
                # 如果行没有足够的列，用空单元格填充
                cells.append(" ")
        
        # 添加行到Markdown表格
        markdown_table.append("| " + " | ".join(cells) + " |")
        
        # 在第一行后添加分隔行
        if i == 0:
            separator = "| " + " | ".join(["---"] * max_cols) + " |"
            markdown_table.append(separator)
    
    return "\n".join(markdown_table)

def word_to_markdown_advanced(word_file_path):
    """高级Word到Markdown转换 - 使用多种策略"""
    try:
        # 策略1: 使用pandoc转换
        pandoc_result = None
        try:
            pypandoc.get_pandoc_version()
            pandoc_result = pypandoc.convert_file(
                word_file_path, 
                'markdown',
                extra_args=[
                    '--wrap=none',
                    '--extract-media=./media',
                    '--standalone',
                    '--markdown-headings=atx',
                    '--preserve-tabs',
                    '--tab-stop=4'
                ]
            )
        except Exception as e:
            st.info(f"Pandoc转换失败: {str(e)}")
        
        # 策略2: 使用直接解析
        direct_result = word_to_markdown_direct(word_file_path)
        
        # 策略3: 使用python-docx的简化解析
        simple_result = word_to_markdown_simple(word_file_path)
        
        # 选择最佳结果
        results = []
        if pandoc_result and len(pandoc_result.strip()) > 50:
            results.append(("pandoc", pandoc_result))
        if direct_result and len(direct_result.strip()) > 50:
            results.append(("direct", direct_result))
        if simple_result and len(simple_result.strip()) > 50:
            results.append(("simple", simple_result))
        
        if not results:
            st.error("所有转换方法都失败了")
            return None
        
        # 选择最长的结果（通常内容最完整）
        best_method, best_content = max(results, key=lambda x: len(x[1]))
        st.info(f"使用 {best_method} 方法转换，内容长度: {len(best_content)} 字符")
        
        return best_content
        
    except Exception as e:
        st.error(f"高级转换失败: {str(e)}")
        return None

def word_to_markdown_simple(word_file_path):
    """简化的Word到Markdown转换"""
    try:
        doc = Document(word_file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # 简单的格式检测
                if paragraph.style.name.startswith('Heading 1'):
                    content.append(f"# {text}")
                elif paragraph.style.name.startswith('Heading 2'):
                    content.append(f"## {text}")
                elif paragraph.style.name.startswith('Heading 3'):
                    content.append(f"### {text}")
                elif paragraph.style.name.startswith('Heading 4'):
                    content.append(f"#### {text}")
                elif paragraph.style.name.startswith('Heading 5'):
                    content.append(f"##### {text}")
                elif paragraph.style.name.startswith('Heading 6'):
                    content.append(f"###### {text}")
                else:
                    content.append(text)
        
        return "\n".join(content)
    except Exception as e:
        return None

def word_to_markdown(word_file_path):
    """将Word文档转换为Markdown格式 - 主入口"""
    return word_to_markdown_advanced(word_file_path)

def markdown_to_word_with_fonts(markdown_content, output_path, fonts_info):
    """使用字体信息生成Word文档"""
    try:
        # 首先尝试使用pandoc转换
        try:
            pypandoc.get_pandoc_version()
            
            # 使用优化的pandoc参数来转换
            pypandoc.convert_text(
                markdown_content, 
                'docx', 
                format='markdown',
                outputfile=output_path,
                extra_args=[
                    '--standalone',  # 生成独立文档
                    '--wrap=none',  # 不自动换行
                    '--toc',  # 生成目录
                    '--toc-depth=6',  # 目录深度
                    '--syntax-highlighting=pygments',  # 代码高亮
                    '--metadata', 'title=翻译文档',  # 设置文档标题
                    '--metadata', 'author=Word翻译工具',  # 设置作者
                ]
            )
            
            # 转换成功后，应用字体信息
            if fonts_info:
                apply_fonts_to_document(output_path, fonts_info)
            
            return True
        except Exception as e:
            st.warning(f"Pandoc转换失败: {str(e)}，尝试直接创建...")
            return markdown_to_word_direct(markdown_content, output_path, fonts_info)
            
    except Exception as e:
        st.error(f"Word生成失败: {str(e)}")
        return False

def apply_fonts_to_document(doc_path, fonts_info):
    """将字体信息应用到Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document(doc_path)
        
        # 应用默认字体
        default_font = fonts_info.get('default_font', '宋体')
        default_size = fonts_info.get('default_size', 12)
        
        # 设置文档默认样式
        style = doc.styles['Normal']
        style.font.name = default_font
        style.font.size = Pt(default_size)
        
        # 应用标题字体
        for level in range(1, 7):
            heading_key = f'heading_{level}'
            if heading_key in fonts_info.get('heading_fonts', {}):
                heading_font = fonts_info['heading_fonts'][heading_key]
                try:
                    heading_style = doc.styles[f'Heading {level}']
                    heading_style.font.name = heading_font.get('name', default_font)
                    heading_style.font.size = Pt(heading_font.get('size', default_size))
                except:
                    pass
        
        # 保存文档
        doc.save(doc_path)
        return True
        
    except Exception as e:
        st.warning(f"字体应用失败: {str(e)}")
        return False

def translate_markdown(markdown_content, target_lang='zh', api_key=None):
    """使用OpenAI翻译Markdown内容"""
    try:
        if not api_key:
            st.error("请设置OpenAI API密钥")
            return markdown_content
            
        # 构建翻译提示 - 使用英文prompt获得更好效果
        lang_map = {
            'zh': 'Chinese',
            'en': 'English', 
            'ja': 'Japanese',
            'ko': 'Korean',
            'fr': 'French',
            'de': 'German',
            'es': 'Spanish',
            'ru': 'Russian'
        }
        
        target_lang_name = lang_map.get(target_lang, 'Chinese')
        
        prompt = f"""Please translate all text content in the following Markdown to {target_lang_name} while keeping the Markdown format completely unchanged:

{markdown_content}

Critical translation requirements:
1. Keep all Markdown syntax format completely unchanged (such as #headings, **bold**, *italic*, [links](url), lists, tables, etc.)
2. Only translate text content, do not change any format structure
3. Keep code blocks, mathematical formulas, and URL links unchanged
4. Keep all punctuation and special characters in their original positions
5. Ensure the translated document format is completely consistent with the original document
6. Unify all non-{target_lang_name} text to be translated to {target_lang_name}"""
        
        # 设置API密钥
        openai.api_key = api_key
        
        # 使用兼容的OpenAI调用方式
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": f"You are a professional document translation assistant, specialized in translating all text content in documents to {target_lang_name} while completely preserving the original document's format structure. You must ensure that the translated document format is completely consistent with the original document, including all Markdown syntax, punctuation, and special character positions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.1  # 降低温度以获得更一致的格式保持
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"翻译失败: {str(e)}")
        return markdown_content

def markdown_to_word(markdown_content, output_path):
    """将Markdown转换为Word文档 - 优化版，模拟Try pandoc!界面功能"""
    try:
        # 确保pandoc可用
        try:
            pypandoc.get_pandoc_version()
        except OSError:
            st.info("正在下载pandoc...")
            pypandoc.download_pandoc()
        
        # 使用优化的pandoc参数来转换 - 参考Try pandoc!最佳实践
        pypandoc.convert_text(
            markdown_content, 
            'docx', 
            format='markdown',
            outputfile=output_path,
            extra_args=[
                '--standalone',  # 生成独立文档
                '--wrap=none',  # 不自动换行
                '--toc',  # 生成目录
                '--toc-depth=6',  # 目录深度
                '--number-sections',  # 自动编号章节
                '--syntax-highlighting=pygments',  # 代码高亮
                '--metadata', 'title=翻译文档',  # 设置文档标题
                '--metadata', 'author=Word翻译工具',  # 设置作者
                '--embed-resources',  # 嵌入资源
                '--citeproc',  # 处理引用
                '--mathjax',  # 数学公式支持
            ]
        )
        return True
    except Exception as e:
        st.warning(f"Pandoc转换失败: {str(e)}，尝试备用方法...")
        return markdown_to_word_direct(markdown_content, output_path)

def markdown_to_word_advanced(markdown_content, output_path, options=None):
    """高级Markdown到Word转换 - 支持自定义选项，模拟Try pandoc!界面"""
    try:
        # 确保pandoc可用
        try:
            pypandoc.get_pandoc_version()
        except OSError:
            st.info("正在下载pandoc...")
            pypandoc.download_pandoc()
        
        # 默认选项
        default_options = {
            'standalone': True,
            'toc': True,
            'number_sections': True,
            'embed_resources': True,
            'citeproc': False,
            'mathjax': False,
            'syntax_highlighting': 'pygments',
            'wrap': 'none',
            'toc_depth': 6
        }
        
        # 合并用户选项
        if options:
            default_options.update(options)
        
        # 构建参数列表
        extra_args = []
        
        if default_options.get('standalone'):
            extra_args.append('--standalone')
        
        if default_options.get('toc'):
            extra_args.append('--toc')
            extra_args.append(f"--toc-depth={default_options.get('toc_depth', 6)}")
        
        if default_options.get('number_sections'):
            extra_args.append('--number-sections')
        
        if default_options.get('embed_resources'):
            extra_args.append('--embed-resources')
        
        if default_options.get('citeproc'):
            extra_args.append('--citeproc')
        
        if default_options.get('mathjax'):
            extra_args.append('--mathjax')
        
        if default_options.get('syntax_highlighting'):
            extra_args.append(f"--syntax-highlighting={default_options.get('syntax_highlighting')}")
        
        if default_options.get('wrap'):
            extra_args.append(f"--wrap={default_options.get('wrap')}")
        
        # 添加元数据
        extra_args.extend(['--metadata', 'title=翻译文档'])
        extra_args.extend(['--metadata', 'author=Word翻译工具'])
        
        # 执行转换
        pypandoc.convert_text(
            markdown_content, 
            'docx', 
            format='markdown',
            outputfile=output_path,
            extra_args=extra_args
        )
        return True
        
    except Exception as e:
        st.warning(f"高级Pandoc转换失败: {str(e)}，尝试备用方法...")
        return markdown_to_word_direct(markdown_content, output_path)

def markdown_to_word_direct(markdown_content, output_path, fonts_info=None):
    """直接使用python-docx创建Word文档 - 优化版"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        doc = Document()
        
        # 使用提取的字体信息或默认字体
        if fonts_info:
            default_font = fonts_info.get('default_font', '宋体')
            default_size = fonts_info.get('default_size', 12)
        else:
            default_font = '宋体'
            default_size = 12
        
        # 设置文档样式
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = default_font
            font.size = Pt(default_size)
        except Exception as e:
            print(f"设置文档样式失败: {str(e)}")
        
        # 按行分割内容
        lines = markdown_content.split('\n')
        in_table = False
        table_data = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # 跳过空行
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue
            
            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                
                # 清理标题文本
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 移除粗体标记
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # 移除斜体标记
                
                if level == 1:
                    heading = doc.add_heading(text, level=1)
                elif level == 2:
                    heading = doc.add_heading(text, level=2)
                elif level == 3:
                    heading = doc.add_heading(text, level=3)
                elif level == 4:
                    heading = doc.add_heading(text, level=4)
                elif level == 5:
                    heading = doc.add_heading(text, level=5)
                else:
                    heading = doc.add_heading(text, level=6)
                
                # 设置标题字体（使用原文档的字体）
                if fonts_info and f'heading_{level}' in fonts_info.get('heading_fonts', {}):
                    heading_font = fonts_info['heading_fonts'][f'heading_{level}']
                    heading.font.name = heading_font.get('name', default_font)
                    heading.font.size = Pt(heading_font.get('size', default_size))
                else:
                    heading.font.name = default_font
                    heading.font.size = Pt(default_size)
                
                # 设置标题样式
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                in_table = False
                table_data = []
            
            # 处理表格 - 改进版
            elif '|' in line and line.count('|') >= 2:
                if not in_table:
                    in_table = True
                    table_data = []
                
                # 解析表格行，改进处理逻辑
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # 检查是否是分隔行（如 |---|---|）
                is_separator = all(
                    cell.replace('-', '').replace(' ', '').replace(':', '') == '' 
                    for cell in cells
                )
                
                # 如果不是分隔行且有内容，添加到表格数据
                if not is_separator and cells:
                    # 确保所有单元格都有内容，空单元格用空格填充
                    normalized_cells = []
                    for cell in cells:
                        if cell.strip():
                            normalized_cells.append(cell.strip())
                        else:
                            normalized_cells.append(" ")
                    table_data.append(normalized_cells)
            
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                # 清理列表文本
                text = _clean_markdown_text(text)
                doc.add_paragraph(text, style='List Bullet')
                in_table = False
                table_data = []
            
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line).strip()
                text = _clean_markdown_text(text)
                doc.add_paragraph(text, style='List Number')
                in_table = False
                table_data = []
            
            # 处理普通段落
            else:
                # 如果之前在处理表格，先完成表格
                if in_table and table_data:
                    _create_table_from_data(doc, table_data)
                    in_table = False
                    table_data = []
                
                # 处理普通段落
                text = _clean_markdown_text(line)
                if text:  # 只添加非空文本
                    paragraph = doc.add_paragraph()
                    try:
                        _add_formatted_text_advanced(paragraph, line, fonts_info)
                    except Exception as e:
                        # 如果格式化失败，添加普通文本
                        paragraph.add_run(text)
        
        # 处理最后的表格
        if in_table and table_data:
            _create_table_from_data(doc, table_data)
        
        # 保存文档
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"直接创建Word文档失败: {str(e)}")
        return False

def _clean_markdown_text(text):
    """清理Markdown文本，移除多余的格式标记"""
    import re
    
    # 移除多余的星号
    text = re.sub(r'\*{3,}', '**', text)  # 将3个或更多星号替换为2个
    text = re.sub(r'\*{2,}', '**', text)  # 将2个或更多星号替换为2个
    
    # 移除多余的换行
    text = re.sub(r'\n+', ' ', text)
    
    return text.strip()

def _validate_and_fix_table_data(table_data):
    """验证和修复表格数据，解决信息错位和消失问题"""
    if not table_data:
        return []
    
    # 获取最大列数
    max_cols = max(len(row) for row in table_data) if table_data else 0
    if max_cols == 0:
        return []
    
    # 标准化所有行
    fixed_data = []
    for row in table_data:
        fixed_row = []
        for j in range(max_cols):
            if j < len(row) and row[j]:
                cell_text = str(row[j]).strip()
                fixed_row.append(cell_text if cell_text else " ")
            else:
                fixed_row.append(" ")
        fixed_data.append(fixed_row)
    
    return fixed_data

def _create_table_from_data(doc, table_data):
    """从数据创建表格 - 改进版，解决信息错位和消失问题"""
    if not table_data:
        return
    
    # 验证和修复表格数据
    fixed_data = _validate_and_fix_table_data(table_data)
    if not fixed_data:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(fixed_data), cols=len(fixed_data[0]))
    table.style = 'Table Grid'
    
    # 填充数据，确保所有单元格都被正确填充
    for i, row_data in enumerate(fixed_data):
        for j, cell_data in enumerate(row_data):
            if i < len(table.rows) and j < len(table.rows[i].cells):
                # 设置单元格文本
                table.rows[i].cells[j].text = cell_data
                
                # 设置单元格样式
                cell = table.rows[i].cells[j]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1  # 居中对齐
                
                # 设置字体
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)

def _add_formatted_text_advanced(paragraph, text, fonts_info=None):
    """高级格式化文本添加"""
    import re
    from docx.shared import Pt
    
    # 获取字体信息
    if fonts_info:
        default_font = fonts_info.get('default_font', '宋体')
        default_size = fonts_info.get('default_size', 12)
        bold_fonts = fonts_info.get('bold_fonts', [])
        italic_fonts = fonts_info.get('italic_fonts', [])
    else:
        default_font = '宋体'
        default_size = 12
        bold_fonts = []
        italic_fonts = []
    
    try:
        # 分割文本，保持格式标记
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if not part:
                continue
            elif part.startswith('**') and part.endswith('**'):
                # 粗体
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                # 使用粗体字体
                if bold_fonts:
                    run.font.name = bold_fonts[0]
                else:
                    run.font.name = default_font
                run.font.size = Pt(default_size)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # 斜体
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                # 使用斜体字体
                if italic_fonts:
                    run.font.name = italic_fonts[0]
                else:
                    run.font.name = default_font
                run.font.size = Pt(default_size)
            else:
                # 普通文本
                run = paragraph.add_run(part)
                run.font.name = default_font
                run.font.size = Pt(default_size)
    except Exception as e:
        # 如果格式化失败，添加普通文本
        paragraph.add_run(text)

def _add_formatted_text(paragraph, text):
    """添加格式化文本到段落"""
    import re
    
    # 处理粗体 **text**
    bold_pattern = r'\*\*(.*?)\*\*'
    italic_pattern = r'\*(.*?)\*'
    
    # 分割文本并处理格式
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if not part:
            continue
        elif part.startswith('**') and part.endswith('**'):
            # 粗体
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # 斜体
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # 普通文本
            paragraph.add_run(part)

def main():
    st.title("📄 Word文档翻译工具")
    st.markdown("**统一语言翻译**: 将文档中的所有语言统一翻译为您选择的目标语言，完全保持原文档格式")
    
    # API密钥设置
    st.sidebar.subheader("🔑 OpenAI设置")
    api_key = st.sidebar.text_input(
        "OpenAI API密钥",
        type="password",
        help="请输入您的OpenAI API密钥"
    )
    
    if not api_key:
        st.warning("⚠️ 请先在侧边栏设置OpenAI API密钥")
        st.stop()
    
    # 高级转换选项
    st.sidebar.subheader("⚙️ 高级转换选项")
    
    # 转换方法选择
    conversion_method = st.sidebar.selectbox(
        "转换方法",
        ["自动选择", "Pandoc优化", "Pandoc高级", "直接转换"],
        help="选择Markdown到Word的转换方法"
    )
    
    # 高级Pandoc选项
    if conversion_method in ["Pandoc高级"]:
        st.sidebar.subheader("🔧 Pandoc高级选项")
        
        standalone = st.sidebar.checkbox("独立文档", value=True, help="生成独立文档")
        toc = st.sidebar.checkbox("生成目录", value=True, help="自动生成目录")
        number_sections = st.sidebar.checkbox("章节编号", value=True, help="自动编号章节")
        embed_resources = st.sidebar.checkbox("嵌入资源", value=True, help="嵌入所有资源")
        citeproc = st.sidebar.checkbox("处理引用", value=False, help="处理学术引用")
        mathjax = st.sidebar.checkbox("数学公式", value=False, help="支持数学公式")
        
        syntax_highlighting = st.sidebar.selectbox(
            "代码高亮",
            ["pygments", "kate", "breezedark", "none"],
            index=0,
            help="选择代码高亮样式"
        )
        
        wrap_mode = st.sidebar.selectbox(
            "文本换行",
            ["none", "auto", "preserve"],
            index=0,
            help="选择文本换行模式"
        )
        
        toc_depth = st.sidebar.slider(
            "目录深度",
            min_value=1,
            max_value=6,
            value=6,
            help="目录的最大深度"
        )
        
        # 构建高级选项
        advanced_options = {
            'standalone': standalone,
            'toc': toc,
            'number_sections': number_sections,
            'embed_resources': embed_resources,
            'citeproc': citeproc,
            'mathjax': mathjax,
            'syntax_highlighting': syntax_highlighting,
            'wrap': wrap_mode,
            'toc_depth': toc_depth
        }
    else:
        advanced_options = None
    
    # 语言选择
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("目标语言")
        target_languages = {
            '中文': 'zh',
            '英文': 'en',
            '日文': 'ja',
            '韩文': 'ko',
            '法文': 'fr',
            '德文': 'de',
            '西班牙文': 'es',
            '俄文': 'ru'
        }
        
        selected_lang = st.selectbox(
            "选择统一翻译的目标语言",
            options=list(target_languages.keys()),
            index=0,
            help="文档中的所有其他语言将被统一翻译为此语言"
        )
        target_lang_code = target_languages[selected_lang]
        
        st.info(f"💡 选择{selected_lang}后，文档中的所有其他语言文本将被统一翻译为{selected_lang}，格式完全保持不变")
    
    with col2:
        st.subheader("文件上传")
        uploaded_file = st.file_uploader(
            "选择Word文档",
            type=['docx'],
            help="请上传.docx格式的Word文档"
        )
    
    if uploaded_file is not None:
        # 保存上传的文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            # 步骤0: 提取字体信息
            st.subheader("步骤0: 提取文档字体信息")
            with st.spinner("正在分析原文档字体..."):
                fonts_info = extract_document_fonts(tmp_file_path)
            
            # 显示字体信息
            if fonts_info:
                st.success("✅ 字体信息提取成功")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("默认字体", fonts_info.get('default_font', '宋体'))
                with col2:
                    st.metric("默认字号", f"{fonts_info.get('default_size', 12)}pt")
                with col3:
                    st.metric("标题样式", f"{len(fonts_info.get('heading_fonts', {}))}种")
                
                # 显示详细字体信息
                with st.expander("查看详细字体信息", expanded=False):
                    st.json(fonts_info)
            
            # 步骤1: Word转Markdown
            st.subheader("步骤1: 转换为Markdown格式")
            with st.spinner("正在转换Word文档为Markdown格式..."):
                markdown_content = word_to_markdown(tmp_file_path)
            
            if markdown_content:
                # 分析Markdown质量
                content_length = len(markdown_content.strip())
                heading_count = markdown_content.count('#')
                paragraph_count = len([p for p in markdown_content.split('\n') if p.strip() and not p.startswith('#')])
                
                st.success("✅ Word文档已成功转换为Markdown格式")
                
                # 显示转换质量信息
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("内容长度", f"{content_length} 字符")
                with col2:
                    st.metric("标题数量", f"{heading_count} 个")
                with col3:
                    st.metric("段落数量", f"{paragraph_count} 个")
                
                # 显示Markdown预览
                with st.expander("查看Markdown内容", expanded=False):
                    st.markdown(markdown_content)
                    
                # 显示转换质量评估
                if content_length < 100:
                    st.warning("⚠️ 转换内容较少，可能格式识别不完整")
                elif heading_count == 0 and paragraph_count < 3:
                    st.warning("⚠️ 未检测到明显的标题结构，可能影响翻译质量")
                else:
                    st.success("✅ Markdown转换质量良好")
                
                # 步骤2: 翻译
                st.subheader("步骤2: 统一语言翻译")
                with st.spinner(f"正在将文档中的所有语言统一翻译为{selected_lang}..."):
                    translated_content = translate_markdown(markdown_content, target_lang_code, api_key)
                
                if translated_content:
                    st.success(f"✅ 文档已成功统一翻译为{selected_lang}，格式完全保持不变")
                    
                    # 显示翻译后的Markdown预览
                    with st.expander("查看统一翻译后的Markdown内容", expanded=False):
                        st.markdown(translated_content)
                    
                    # 步骤3: 生成Word文档
                    st.subheader("步骤3: 生成翻译后的Word文档")
                    
                    # 创建下载按钮
                    output_filename = f"translated_{uploaded_file.name}"
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as output_file:
                        output_path = output_file.name
                    
                    with st.spinner("正在生成Word文档..."):
                        # 根据用户选择的方法进行转换
                        success = False
                        
                        if conversion_method == "自动选择":
                            # 优先使用字体保持的转换
                            success = markdown_to_word_with_fonts(translated_content, output_path, fonts_info)
                            if not success:
                                # 备用方案
                                success = markdown_to_word(translated_content, output_path)
                        
                        elif conversion_method == "Pandoc优化":
                            success = markdown_to_word(translated_content, output_path)
                        
                        elif conversion_method == "Pandoc高级":
                            success = markdown_to_word_advanced(translated_content, output_path, advanced_options)
                        
                        elif conversion_method == "直接转换":
                            success = markdown_to_word_direct(translated_content, output_path, fonts_info)
                        
                        if success:
                            st.success(f"✅ 翻译后的Word文档已生成（使用{conversion_method}方法）")
                            
                            # 格式对比分析
                            st.subheader("📊 格式对比分析")
                            
                            # 分析Markdown格式
                            md_headings = translated_content.count('#')
                            md_lists = translated_content.count('- ') + translated_content.count('* ')
                            md_bold = translated_content.count('**')
                            md_italic = translated_content.count('*') - md_bold
                            md_tables = translated_content.count('|') // 3  # 估算表格数量
                            
                            col1, col2, col3, col4, col5 = st.columns(5)
                            with col1:
                                st.metric("标题数量", md_headings)
                            with col2:
                                st.metric("列表项", md_lists)
                            with col3:
                                st.metric("粗体标记", md_bold)
                            with col4:
                                st.metric("斜体标记", md_italic)
                            with col5:
                                st.metric("表格数量", md_tables)
                            
                            # 显示格式保持说明
                            st.info("""
                            **格式保持说明**:
                            - ✅ 标题结构已转换为Word标题样式
                            - ✅ 列表已转换为Word列表格式
                            - ✅ 粗体/斜体格式已保持
                            - ✅ 段落结构已保持
                            - ✅ 表格结构已优化，解决信息错位和消失问题
                            - 📝 如发现格式差异，请检查Word文档的样式设置
                            """)
                            
                            # 读取生成的文件并提供下载
                            with open(output_path, 'rb') as f:
                                file_data = f.read()
                            
                            st.download_button(
                                label="📥 下载翻译后的Word文档",
                                data=file_data,
                                file_name=output_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error("❌ 生成Word文档失败")
            
        except Exception as e:
            st.error(f"处理过程中出现错误: {str(e)}")
        
        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_file_path)
                if 'output_path' in locals():
                    os.unlink(output_path)
            except:
                pass
    
    # 使用说明
    st.markdown("---")
    st.subheader("📖 使用说明")
    st.markdown("""
    ### 🎯 核心功能
    **统一语言翻译**: 将文档中的所有语言文本统一翻译为您选择的目标语言，完全保持原文档格式
    
    ### 📋 使用步骤
    1. **设置API密钥**: 在侧边栏输入您的OpenAI API密钥
    2. **上传文档**: 选择要翻译的Word文档（.docx格式）
    3. **选择目标语言**: 从下拉菜单中选择统一翻译的目标语言
    4. **自动处理**: 系统会自动将Word转换为Markdown，统一翻译所有语言，然后生成新的Word文档
    5. **下载结果**: 点击下载按钮获取统一翻译后的Word文档
    
    ### ✨ 特色功能
    - **格式完全保持**: 所有Markdown语法、标点符号、特殊字符位置完全不变
    - **统一语言**: 文档中的所有其他语言将被统一翻译为目标语言
    - **高质量翻译**: 使用OpenAI GPT-3.5-turbo模型
    - **智能识别**: 自动识别并翻译所有非目标语言的文本
    
    **注意事项**:
    - 支持.docx格式的Word文档
    - 需要有效的OpenAI API密钥
    - 翻译质量取决于OpenAI GPT模型
    - 格式保持度极高，几乎与原文档完全一致
    - 建议文档不要过大以确保处理速度
    """)

if __name__ == "__main__":
    main()
