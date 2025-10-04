"""
简化段落查看器 - 避免复杂状态管理
"""

import streamlit as st
from docx import Document
from typing import List, Dict, Any
import tempfile
import os

class SimpleParagraphViewer:
    """简化段落查看器 - 专注于基本功能"""
    
    def __init__(self):
        self.original_doc = None
        self.translated_doc = None
        self.current_page = 1
        self.total_pages = 0
        self.paragraphs_data = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            self.original_doc = Document(original_path)
            self.translated_doc = Document(translated_path)
            
            # 解析文档结构
            self._parse_document_structure()
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def _parse_document_structure(self):
        """解析文档结构，提取段落信息"""
        self.paragraphs_data = []
        
        # 获取原文档段落
        original_paragraphs = [p.text for p in self.original_doc.paragraphs if p.text.strip()]
        translated_paragraphs = [p.text for p in self.translated_doc.paragraphs if p.text.strip()]
        
        # 按页面组织段落（简单估算）
        paragraphs_per_page = 8  # 每页约8个段落
        
        for i, (orig_text, trans_text) in enumerate(zip(original_paragraphs, translated_paragraphs)):
            page_num = (i // paragraphs_per_page) + 1
            
            self.paragraphs_data.append({
                'index': i,
                'page': page_num,
                'original_text': orig_text,
                'translated_text': trans_text,
                'is_heading': self._is_heading(orig_text),
                'word_count': len(orig_text.split())
            })
        
        self.total_pages = max([p['page'] for p in self.paragraphs_data]) if self.paragraphs_data else 1
    
    def _is_heading(self, text: str) -> bool:
        """判断是否为标题"""
        return (text.isupper() or 
                text.startswith('#') or 
                text.startswith('第') or 
                len(text) < 50 and not text.endswith('.'))
    
    def display_simple_viewer(self):
        """显示简化查看器"""
        if not self.paragraphs_data:
            st.warning("⚠️ 请先加载文档")
            return
        
        # 页面导航
        self._display_simple_navigation()
        
        # 文档内容展示
        self._display_simple_content()
    
    def _display_simple_navigation(self):
        """显示简化页面导航"""
        st.markdown("---")
        
        # 使用session_state来保持页面状态
        if 'simple_current_page' not in st.session_state:
            st.session_state.simple_current_page = 1
        
        self.current_page = st.session_state.simple_current_page
        
        col1, col2, col3, col4 = st.columns([1, 2, 1, 1])
        
        with col1:
            if st.button("⬅️ 上一页", disabled=(self.current_page <= 1), key="simple_prev_page"):
                if self.current_page > 1:
                    st.session_state.simple_current_page = self.current_page - 1
                    st.rerun()
        
        with col2:
            st.markdown(f"**第 {self.current_page} 页 / 共 {self.total_pages} 页**")
        
        with col3:
            if st.button("下一页 ➡️", disabled=(self.current_page >= self.total_pages), key="simple_next_page"):
                if self.current_page < self.total_pages:
                    st.session_state.simple_current_page = self.current_page + 1
                    st.rerun()
        
        with col4:
            # 页面跳转
            target_page = st.number_input(
                "跳转到", 
                min_value=1, 
                max_value=self.total_pages, 
                value=self.current_page,
                key="simple_page_jumper"
            )
            if target_page != self.current_page:
                st.session_state.simple_current_page = target_page
                st.rerun()
    
    def _display_simple_content(self):
        """显示简化内容"""
        # 获取当前页面的段落
        current_paragraphs = [p for p in self.paragraphs_data if p['page'] == self.current_page]
        
        if not current_paragraphs:
            st.info("📄 该页面没有内容")
            return
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📝 原文")
            for i, para in enumerate(current_paragraphs):
                if para['is_heading']:
                    st.markdown(f"#### {para['original_text']}")
                else:
                    st.markdown(f"**段落 {para['index']+1}:**")
                    st.text(para['original_text'])
                st.markdown("---")
        
        with col2:
            st.markdown("### 🌐 译文")
            for i, para in enumerate(current_paragraphs):
                if para['is_heading']:
                    st.markdown(f"#### {para['translated_text']}")
                else:
                    st.markdown(f"**段落 {para['index']+1}:**")
                    st.text(para['translated_text'])
                st.markdown("---")
    
    def get_document_summary(self):
        """获取文档摘要"""
        if not self.paragraphs_data:
            return {}
        
        total_paragraphs = len(self.paragraphs_data)
        total_pages = self.total_pages
        headings = len([p for p in self.paragraphs_data if p['is_heading']])
        avg_words = sum(p['word_count'] for p in self.paragraphs_data) / total_paragraphs
        
        return {
            'total_paragraphs': total_paragraphs,
            'total_pages': total_pages,
            'headings': headings,
            'avg_words_per_paragraph': round(avg_words, 1)
        }
    
    def display_document_summary(self):
        """显示文档摘要"""
        summary = self.get_document_summary()
        
        if summary:
            st.markdown("### 📊 文档摘要")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("总段落数", summary['total_paragraphs'])
            
            with col2:
                st.metric("总页数", summary['total_pages'])
            
            with col3:
                st.metric("标题数", summary['headings'])
            
            with col4:
                st.metric("平均字数", summary['avg_words_per_paragraph'])
