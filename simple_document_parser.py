"""
简化文档解析器 - 作为备用方案
"""

import streamlit as st
from docx import Document
from typing import Dict, List, Any
import tempfile
import os

class SimpleDocumentParser:
    """简化文档解析器 - 专注于基本功能"""
    
    def __init__(self):
        pass
    
    def parse_document(self, doc_path: str) -> Dict[str, Any]:
        """简化文档解析"""
        try:
            doc = Document(doc_path)
            
            result = {
                'content_layer': [],
                'format_layer': [],
                'layout_layer': [],
                'anchors': {},
                'metadata': {
                    'total_paragraphs': 0,
                    'total_tables': 0,
                    'total_images': 0,
                    'total_pages': 0
                }
            }
            
            # 简化段落解析
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    content_info = {
                        'id': f'para_{i}',
                        'text': paragraph.text.strip(),
                        'type': 'paragraph'
                    }
                    result['content_layer'].append(content_info)
                    
                    # 简化的格式信息
                    format_info = {
                        'id': f'para_{i}',
                        'style': 'Normal',
                        'alignment': 'None',
                        'runs': [{'text': paragraph.text, 'bold': False, 'italic': False, 'underline': False, 'font_name': 'Calibri', 'font_size': 11}]
                    }
                    result['format_layer'].append(format_info)
                    
                    # 简化的布局信息
                    layout_info = {
                        'id': f'para_{i}',
                        'is_heading': False,
                        'heading_level': 0,
                        'page_break_before': False,
                        'page_number': 1
                    }
                    result['layout_layer'].append(layout_info)
                    
                    result['metadata']['total_paragraphs'] += 1
            
            # 简化表格解析
            for i, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            if cell.text.strip():
                                cell_id = f"table_{i}_row_{row_idx}_col_{col_idx}"
                                
                                content_info = {
                                    'id': cell_id,
                                    'text': cell.text.strip(),
                                    'type': 'table_cell',
                                    'table_index': i,
                                    'row': row_idx,
                                    'col': col_idx
                                }
                                result['content_layer'].append(content_info)
                                
                                format_info = {
                                    'id': cell_id,
                                    'style': 'Normal',
                                    'alignment': 'None',
                                    'runs': [{'text': cell.text, 'bold': False, 'italic': False, 'underline': False, 'font_name': 'Calibri', 'font_size': 11}]
                                }
                                result['format_layer'].append(format_info)
                                
                                layout_info = {
                                    'id': cell_id,
                                    'is_heading': False,
                                    'heading_level': 0,
                                    'page_break_before': False,
                                    'page_number': 1
                                }
                                result['layout_layer'].append(layout_info)
                    
                    result['metadata']['total_tables'] += 1
                except Exception as e:
                    print(f"表格 {i} 解析失败: {str(e)}")
                    continue
            
            return result
            
        except Exception as e:
            st.error(f"简化文档解析失败: {str(e)}")
            return None
    
    def create_simple_translated_document(self, original_path: str, translated_content: List[Dict], output_path: str) -> bool:
        """创建简化的翻译文档"""
        try:
            # 读取原文档
            doc = Document(original_path)
            
            # 创建翻译映射
            translation_map = {}
            for item in translated_content:
                if 'text' in item:
                    translation_map[item['id']] = item['text']
            
            # 更新段落
            para_index = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_id = f'para_{para_index}'
                    if para_id in translation_map:
                        paragraph.text = translation_map[para_id]
                    para_index += 1
            
            # 更新表格
            table_index = 0
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            cell_id = f"table_{table_index}_row_{row_idx}_col_{col_idx}"
                            if cell_id in translation_map:
                                cell.text = translation_map[cell_id]
                table_index += 1
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            st.error(f"创建翻译文档失败: {str(e)}")
            return False
