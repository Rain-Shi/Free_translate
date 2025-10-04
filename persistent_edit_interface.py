"""
持久化编辑界面 - 完全避免跳回主页
"""

import streamlit as st
from docx import Document
from typing import List, Dict, Any
import tempfile
import os

class PersistentEditInterface:
    """持久化编辑界面 - 使用session_state完全避免跳回"""
    
    def __init__(self):
        self.original_paragraphs = []
        self.translated_paragraphs = []
        self.edited_paragraphs = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            # 读取原文档
            original_doc = Document(original_path)
            self.original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
            
            # 读取翻译文档
            translated_doc = Document(translated_path)
            self.translated_paragraphs = [p.text.strip() for p in translated_doc.paragraphs if p.text.strip()]
            
            # 初始化编辑后的段落
            self.edited_paragraphs = self.translated_paragraphs.copy()
            
            # 初始化session_state
            if 'edit_interface_initialized' not in st.session_state:
                st.session_state.edit_interface_initialized = True
                for i in range(len(self.edited_paragraphs)):
                    st.session_state[f"edited_text_{i}"] = self.edited_paragraphs[i]
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def display_persistent_edit_interface(self):
        """显示持久化编辑界面"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            st.warning("⚠️ 请先加载文档")
            return
        
        st.markdown("---")
        st.subheader("📝 持久化编辑界面")
        
        # 显示统计信息
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("原文档段落数", len(self.original_paragraphs))
        
        with col2:
            st.metric("译文段落数", len(self.translated_paragraphs))
        
        with col3:
            st.metric("可编辑段落数", len(self.edited_paragraphs))
        
        # 使用容器来避免重新运行
        with st.container():
            # 一次性显示所有段落的左右编辑界面
            max_paragraphs = min(len(self.original_paragraphs), len(self.translated_paragraphs))
            
            if max_paragraphs > 0:
                # 显示所有段落的左右编辑界面
                for i in range(max_paragraphs):
                    self._display_paragraph_edit_persistent(i)
        
        # 操作按钮区域
        st.markdown("---")
        st.subheader("🔧 操作按钮")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 重置所有为原文", key="reset_all_original"):
                for i in range(len(self.original_paragraphs)):
                    st.session_state[f"edited_text_{i}"] = self.original_paragraphs[i]
                    self.edited_paragraphs[i] = self.original_paragraphs[i]
                st.success("✅ 已重置所有段落为原文")
                st.rerun()
        
        with col2:
            if st.button("🔄 重置所有为译文", key="reset_all_translated"):
                for i in range(len(self.translated_paragraphs)):
                    st.session_state[f"edited_text_{i}"] = self.translated_paragraphs[i]
                    self.edited_paragraphs[i] = self.translated_paragraphs[i]
                st.success("✅ 已重置所有段落为译文")
                st.rerun()
        
        with col3:
            if st.button("📋 复制所有原文", key="copy_all_original"):
                for i in range(len(self.original_paragraphs)):
                    st.session_state[f"edited_text_{i}"] = self.original_paragraphs[i]
                    self.edited_paragraphs[i] = self.original_paragraphs[i]
                st.success("✅ 已复制所有原文")
                st.rerun()
    
    def _display_paragraph_edit_persistent(self, para_index: int):
        """显示单个段落的持久化编辑界面"""
        st.markdown(f"### ✏️ 段落 {para_index + 1}")
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📝 原文**")
            original_text = self.original_paragraphs[para_index] if para_index < len(self.original_paragraphs) else ""
            st.text_area(
                "原文内容",
                value=original_text,
                height=150,
                key=f"original_display_{para_index}",
                disabled=True
            )
            
            # 原文统计
            st.markdown(f"**字数**: {len(original_text)}")
        
        with col2:
            st.markdown("**🌐 译文 (可编辑)**")
            # 使用session_state来保持编辑状态
            edit_key = f"edited_text_{para_index}"
            
            # 确保session_state中有这个键
            if edit_key not in st.session_state:
                st.session_state[edit_key] = self.edited_paragraphs[para_index] if para_index < len(self.edited_paragraphs) else ""
            
            # 使用st.text_input而不是st.text_area，减少重新运行
            edited_text = st.text_input(
                "译文内容 (可编辑)",
                value=st.session_state[edit_key],
                key=f"translated_edit_{para_index}",
                help="您可以在这里编辑译文内容"
            )
            
            # 更新编辑后的段落
            if edited_text != st.session_state[edit_key]:
                st.session_state[edit_key] = edited_text
                self.edited_paragraphs[para_index] = edited_text
            
            # 译文统计
            st.markdown(f"**字数**: {len(edited_text)}")
        
        # 对比统计
        if original_text and edited_text:
            length_ratio = len(edited_text) / len(original_text) if original_text else 1
            if abs(length_ratio - 1.0) > 0.1:
                st.markdown(f"**长度比例**: {length_ratio:.2f}")
    
    def create_final_document(self, output_path: str):
        """创建最终文档"""
        try:
            # 读取原文档结构
            original_doc = Document()
            
            # 创建新文档
            new_doc = Document()
            
            # 添加标题
            new_doc.add_heading('翻译后的文档', 0)
            
            # 添加编辑后的段落
            for i, edited_text in enumerate(self.edited_paragraphs):
                if edited_text.strip():
                    new_doc.add_paragraph(edited_text)
            
            # 保存文档
            new_doc.save(output_path)
            return True
        except Exception as e:
            st.error(f"❌ 创建最终文档失败: {str(e)}")
            return False
    
    def get_edit_summary(self):
        """获取编辑摘要"""
        if not self.original_paragraphs or not self.edited_paragraphs:
            return {}
        
        # 计算修改统计
        modified_count = 0
        for i in range(min(len(self.translated_paragraphs), len(self.edited_paragraphs))):
            if self.translated_paragraphs[i] != self.edited_paragraphs[i]:
                modified_count += 1
        
        return {
            'total_paragraphs': len(self.edited_paragraphs),
            'modified_paragraphs': modified_count,
            'modification_rate': modified_count / len(self.edited_paragraphs) if self.edited_paragraphs else 0
        }
    
    def display_edit_summary(self):
        """显示编辑摘要"""
        summary = self.get_edit_summary()
        
        if summary:
            st.markdown("### 📊 编辑摘要")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("总段落数", summary['total_paragraphs'])
            
            with col2:
                st.metric("已修改段落数", summary['modified_paragraphs'])
            
            with col3:
                st.metric("修改率", f"{summary['modification_rate']:.1%}")
