"""
测试改进的Markdown提取功能
"""
import os
import tempfile
from docx import Document

def create_test_document():
    """创建一个测试Word文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档标题', 0)
    doc.add_heading('第一章 介绍', 1)
    doc.add_heading('1.1 背景', 2)
    
    # 添加段落
    doc.add_paragraph('这是一个测试段落，包含普通文本。')
    doc.add_paragraph('这是另一个段落，包含**粗体文本**和*斜体文本*。')
    
    # 添加列表
    doc.add_paragraph('主要特点：', style='List Bullet')
    doc.add_paragraph('特点一：高质量翻译', style='List Bullet')
    doc.add_paragraph('特点二：格式保持', style='List Bullet')
    
    # 保存到临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name

def test_direct_extraction():
    """测试直接解析方法"""
    print("测试直接解析方法...")
    
    # 创建测试文档
    test_file = create_test_document()
    
    try:
        # 导入我们的函数
        import sys
        sys.path.append('.')
        from app import word_to_markdown_direct
        
        # 测试直接解析
        result = word_to_markdown_direct(test_file)
        
        if result:
            print("直接解析成功")
            print("解析结果:")
            print("-" * 50)
            print(result)
            print("-" * 50)
            
            # 分析结果质量
            content_length = len(result.strip())
            heading_count = result.count('#')
            paragraph_count = len([p for p in result.split('\n') if p.strip() and not p.startswith('#')])
            
            print(f"\n质量分析:")
            print(f"内容长度: {content_length} 字符")
            print(f"标题数量: {heading_count} 个")
            print(f"段落数量: {paragraph_count} 个")
            
            if content_length > 100 and heading_count > 0:
                print("解析质量良好")
            else:
                print("解析质量可能不够理想")
        else:
            print("直接解析失败")
            
    except Exception as e:
        print(f"测试失败: {e}")
    finally:
        # 清理临时文件
        try:
            os.unlink(test_file)
        except:
            pass

if __name__ == "__main__":
    test_direct_extraction()
