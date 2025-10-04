"""
测试Word文档翻译工具的功能
"""
import os
import tempfile
from docx import Document
import pypandoc
import openai

def test_dependencies():
    """测试所有依赖是否正确安装"""
    print("测试依赖包...")
    
    try:
        from docx import Document
        print("python-docx 导入成功")
    except ImportError as e:
        print(f"python-docx 导入失败: {e}")
        return False
    
    try:
        import pypandoc
        version = pypandoc.get_pandoc_version()
        print(f"pypandoc 导入成功，版本: {version}")
    except Exception as e:
        print(f"pypandoc 导入失败: {e}")
        return False
    
    try:
        import openai
        print("openai 导入成功")
    except ImportError as e:
        print(f"openai 导入失败: {e}")
        return False
    
    try:
        import streamlit
        print("streamlit 导入成功")
    except ImportError as e:
        print(f"streamlit 导入失败: {e}")
        return False
    
    return True

def test_pandoc_conversion():
    """测试pandoc转换功能"""
    print("\n测试pandoc转换功能...")
    
    try:
        # 创建一个简单的测试Word文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试段落。')
        doc.add_paragraph('包含**粗体**和*斜体*文本。')
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            temp_path = tmp_file.name
        
        # 转换为Markdown
        markdown_content = pypandoc.convert_file(temp_path, 'markdown')
        print("Word转Markdown成功")
        print(f"转换结果预览: {markdown_content[:100]}...")
        
        # 清理临时文件
        os.unlink(temp_path)
        
        return True
        
    except Exception as e:
        print(f"pandoc转换测试失败: {e}")
        return False

def main():
    """主测试函数"""
    print("开始测试Word文档翻译工具...")
    
    # 测试依赖
    if not test_dependencies():
        print("\n依赖测试失败，请检查安装")
        return
    
    # 测试pandoc转换
    if not test_pandoc_conversion():
        print("\npandoc转换测试失败")
        return
    
    print("\n所有测试通过！应用已准备就绪。")
    print("使用说明:")
    print("1. 在浏览器中打开 http://localhost:8501")
    print("2. 在侧边栏输入OpenAI API密钥")
    print("3. 上传Word文档并选择翻译语言")
    print("4. 下载翻译后的文档")

if __name__ == "__main__":
    main()
