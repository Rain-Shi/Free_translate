"""
测试文档解析功能
"""

import streamlit as st
from docx import Document
import tempfile
import os
from smart_translator import StructuralParser
from simple_document_parser import SimpleDocumentParser

def create_test_document():
    """创建测试文档"""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_paragraph('Another paragraph with some content.')
    
    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'
    
    # 保存到临时文件
    temp_path = tempfile.mktemp(suffix='.docx')
    doc.save(temp_path)
    return temp_path

def test_parsers():
    """测试解析器"""
    st.title("📄 文档解析测试")
    
    # 创建测试文档
    if st.button("创建测试文档"):
        test_path = create_test_document()
        st.success(f"✅ 测试文档已创建: {test_path}")
        
        # 测试智能解析器
        st.subheader("🧠 智能解析器测试")
        try:
            parser = StructuralParser()
            result = parser.parse_document(test_path)
            
            if result:
                st.success("✅ 智能解析器成功！")
                st.json({
                    'total_paragraphs': result['metadata']['total_paragraphs'],
                    'total_tables': result['metadata']['total_tables'],
                    'content_items': len(result['content_layer'])
                })
            else:
                st.error("❌ 智能解析器失败")
        except Exception as e:
            st.error(f"❌ 智能解析器错误: {str(e)}")
        
        # 测试简化解析器
        st.subheader("🔧 简化解析器测试")
        try:
            simple_parser = SimpleDocumentParser()
            result = simple_parser.parse_document(test_path)
            
            if result:
                st.success("✅ 简化解析器成功！")
                st.json({
                    'total_paragraphs': result['metadata']['total_paragraphs'],
                    'total_tables': result['metadata']['total_tables'],
                    'content_items': len(result['content_layer'])
                })
            else:
                st.error("❌ 简化解析器失败")
        except Exception as e:
            st.error(f"❌ 简化解析器错误: {str(e)}")
        
        # 清理临时文件
        try:
            os.unlink(test_path)
        except:
            pass

if __name__ == "__main__":
    test_parsers()
