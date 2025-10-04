"""
测试文档查看器修复
"""

import streamlit as st
from document_viewer import DocumentViewer
from docx import Document
import tempfile
import os

def create_test_documents():
    """创建测试文档"""
    # 创建原文文档
    original_doc = Document()
    original_doc.add_heading('Test Document', 0)
    original_doc.add_paragraph('This is the first paragraph in English.')
    original_doc.add_paragraph('This is the second paragraph with more content.')
    original_doc.add_heading('Section 1', level=1)
    original_doc.add_paragraph('This is a section paragraph.')
    
    # 创建译文文档
    translated_doc = Document()
    translated_doc.add_heading('测试文档', 0)
    translated_doc.add_paragraph('这是第一段中文内容。')
    translated_doc.add_paragraph('这是第二段中文内容，包含更多信息。')
    translated_doc.add_heading('章节 1', level=1)
    translated_doc.add_paragraph('这是一个章节段落。')
    
    # 保存到临时文件
    original_path = tempfile.mktemp(suffix='.docx')
    translated_path = tempfile.mktemp(suffix='.docx')
    
    original_doc.save(original_path)
    translated_doc.save(translated_path)
    
    return original_path, translated_path

def main():
    """主测试函数"""
    st.set_page_config(
        page_title="文档查看器修复测试",
        page_icon="🔧",
        layout="wide"
    )
    
    st.title("🔧 文档查看器修复测试")
    st.markdown("---")
    
    # 创建测试文档
    if st.button("🚀 创建测试文档并测试查看器"):
        with st.spinner("正在创建测试文档..."):
            original_path, translated_path = create_test_documents()
            
            # 初始化文档查看器
            viewer = DocumentViewer()
            
            # 加载文档
            if viewer.load_documents(original_path, translated_path):
                st.success("✅ 测试文档创建成功！")
                
                # 显示文档摘要
                viewer.display_document_summary()
                
                # 显示文档查看器
                st.markdown("---")
                st.subheader("📖 文档查看器测试")
                viewer.display_document_viewer()
                
                # 清理临时文件
                try:
                    os.unlink(original_path)
                    os.unlink(translated_path)
                except:
                    pass
            else:
                st.error("❌ 测试文档创建失败")
    
    # 修复说明
    st.markdown("---")
    st.subheader("🔧 修复内容")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### ✅ 内容显示修复")
        st.markdown("- 修复译文和原文显示顺序")
        st.markdown("- 正确获取对应版本的文本")
        st.markdown("- 避免内容混淆")
    
    with col2:
        st.markdown("### ✅ 翻页功能修复")
        st.markdown("- 使用session_state保持页面状态")
        st.markdown("- 修复翻页时内容消失问题")
        st.markdown("- 避免跳回主页")
    
    # 测试说明
    st.markdown("---")
    st.subheader("🧪 测试步骤")
    st.markdown("1. 点击'创建测试文档并测试查看器'按钮")
    st.markdown("2. 检查原文和译文是否正确显示")
    st.markdown("3. 测试翻页功能是否正常")
    st.markdown("4. 点击段落按钮测试对比功能")

if __name__ == "__main__":
    main()
