"""
文档查看器 - 支持双版本Word文档展示和段落对比
"""

import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
import os
from typing import List, Dict, Any
import base64

class DocumentViewer:
    """文档查看器 - 支持双版本展示和段落对比"""
    
    def __init__(self):
        self.original_doc = None
        self.translated_doc = None
        self.current_page = 1
        self.total_pages = 0
        self.paragraphs_data = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            self.original_doc = Document(original_path)
            self.translated_doc = Document(translated_path)
            
            # 解析文档结构
            self._parse_document_structure()
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def _parse_document_structure(self):
        """解析文档结构，提取段落信息"""
        self.paragraphs_data = []
        
        # 获取原文档段落
        original_paragraphs = [p.text for p in self.original_doc.paragraphs if p.text.strip()]
        translated_paragraphs = [p.text for p in self.translated_doc.paragraphs if p.text.strip()]
        
        # 按页面组织段落（简单估算）
        paragraphs_per_page = 8  # 每页约8个段落
        
        for i, (orig_text, trans_text) in enumerate(zip(original_paragraphs, translated_paragraphs)):
            page_num = (i // paragraphs_per_page) + 1
            
            self.paragraphs_data.append({
                'index': i,
                'page': page_num,
                'original_text': orig_text,
                'translated_text': trans_text,
                'is_heading': self._is_heading(orig_text),
                'word_count': len(orig_text.split())
            })
        
        self.total_pages = max([p['page'] for p in self.paragraphs_data]) if self.paragraphs_data else 1
    
    def _is_heading(self, text: str) -> bool:
        """判断是否为标题"""
        return (text.isupper() or 
                text.startswith('#') or 
                text.startswith('第') or 
                len(text) < 50 and not text.endswith('.'))
    
    def display_document_viewer(self):
        """显示文档查看器主界面"""
        if not self.paragraphs_data:
            st.warning("⚠️ 请先加载文档")
            return
        
        # 页面导航
        self._display_page_navigation()
        
        # 文档内容展示
        self._display_document_content()
        
        # 段落对比功能
        self._display_paragraph_comparison()
    
    def _display_page_navigation(self):
        """显示页面导航"""
        st.markdown("---")
        
        col1, col2, col3, col4 = st.columns([1, 2, 1, 1])
        
        with col1:
            if st.button("⬅️ 上一页", disabled=(self.current_page <= 1)):
                self.current_page -= 1
                st.rerun()
        
        with col2:
            st.markdown(f"**第 {self.current_page} 页 / 共 {self.total_pages} 页**")
        
        with col3:
            if st.button("下一页 ➡️", disabled=(self.current_page >= self.total_pages)):
                self.current_page += 1
                st.rerun()
        
        with col4:
            # 页面跳转
            target_page = st.number_input(
                "跳转到", 
                min_value=1, 
                max_value=self.total_pages, 
                value=self.current_page,
                key="page_jumper"
            )
            if target_page != self.current_page:
                self.current_page = target_page
                st.rerun()
    
    def _display_document_content(self):
        """显示文档内容"""
        # 获取当前页面的段落
        current_paragraphs = [p for p in self.paragraphs_data if p['page'] == self.current_page]
        
        if not current_paragraphs:
            st.info("📄 该页面没有内容")
            return
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📝 原文")
            self._display_paragraphs(current_paragraphs, "original")
        
        with col2:
            st.markdown("### 🌐 译文")
            self._display_paragraphs(current_paragraphs, "translated")
    
    def _display_paragraphs(self, paragraphs: List[Dict], version: str):
        """显示段落列表"""
        for i, para in enumerate(paragraphs):
            text = para['original_text'] if version == "original" else para['translated_text']
            
            # 段落样式
            if para['is_heading']:
                st.markdown(f"#### {text}")
            else:
                # 可点击的段落
                if st.button(
                    f"📄 段落 {para['index']+1}",
                    key=f"para_{version}_{para['index']}",
                    help=f"点击查看对比 (字数: {para['word_count']})"
                ):
                    # 显示段落对比
                    self._show_paragraph_comparison(para)
            
            # 段落预览（前50个字符）
            preview = text[:50] + "..." if len(text) > 50 else text
            st.text(preview)
            
            st.markdown("---")
    
    def _show_paragraph_comparison(self, paragraph: Dict):
        """显示段落对比"""
        st.markdown("---")
        st.markdown("### 🔍 段落详细对比")
        
        # 创建对比布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📝 原文**")
            st.text_area(
                "原文内容",
                value=paragraph['original_text'],
                height=150,
                key=f"orig_detail_{paragraph['index']}"
            )
        
        with col2:
            st.markdown("**🌐 译文**")
            st.text_area(
                "译文内容", 
                value=paragraph['translated_text'],
                height=150,
                key=f"trans_detail_{paragraph['index']}"
            )
        
        # 统计信息
        self._display_paragraph_stats(paragraph)
        
        # 编辑功能
        self._display_edit_options(paragraph)
    
    def _display_paragraph_stats(self, paragraph: Dict):
        """显示段落统计信息"""
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("原文字数", len(paragraph['original_text']))
        
        with col2:
            st.metric("译文字数", len(paragraph['translated_text']))
        
        with col3:
            ratio = len(paragraph['translated_text']) / len(paragraph['original_text']) if paragraph['original_text'] else 1
            st.metric("长度比例", f"{ratio:.2f}")
        
        with col4:
            st.metric("段落类型", "标题" if paragraph['is_heading'] else "正文")
    
    def _display_edit_options(self, paragraph: Dict):
        """显示编辑选项"""
        st.markdown("### ✏️ 编辑选项")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📝 编辑译文", key=f"edit_{paragraph['index']}"):
                st.session_state[f"editing_{paragraph['index']}"] = True
        
        with col2:
            if st.button("🔄 重新翻译", key=f"retranslate_{paragraph['index']}"):
                st.info("🔄 重新翻译功能开发中...")
        
        with col3:
            if st.button("✅ 确认", key=f"confirm_{paragraph['index']}"):
                st.success("✅ 段落已确认")
    
    def _display_paragraph_comparison(self):
        """显示段落对比功能"""
        st.markdown("---")
        st.markdown("### 🔍 段落对比工具")
        
        # 段落选择器
        paragraph_options = [f"段落 {i+1}" for i in range(len(self.paragraphs_data))]
        selected_para = st.selectbox(
            "选择要对比的段落",
            options=paragraph_options,
            key="paragraph_selector"
        )
        
        if selected_para:
            para_index = int(selected_para.split()[1]) - 1
            if para_index < len(self.paragraphs_data):
                self._show_paragraph_comparison(self.paragraphs_data[para_index])
    
    def get_document_summary(self):
        """获取文档摘要"""
        if not self.paragraphs_data:
            return {}
        
        total_paragraphs = len(self.paragraphs_data)
        total_pages = self.total_pages
        headings = len([p for p in self.paragraphs_data if p['is_heading']])
        avg_words = sum(p['word_count'] for p in self.paragraphs_data) / total_paragraphs
        
        return {
            'total_paragraphs': total_paragraphs,
            'total_pages': total_pages,
            'headings': headings,
            'avg_words_per_paragraph': round(avg_words, 1)
        }
    
    def display_document_summary(self):
        """显示文档摘要"""
        summary = self.get_document_summary()
        
        if summary:
            st.markdown("### 📊 文档摘要")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("总段落数", summary['total_paragraphs'])
            
            with col2:
                st.metric("总页数", summary['total_pages'])
            
            with col3:
                st.metric("标题数", summary['headings'])
            
            with col4:
                st.metric("平均字数", summary['avg_words_per_paragraph'])
