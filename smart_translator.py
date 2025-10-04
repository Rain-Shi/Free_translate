"""
智能文档翻译与格式保真系统
基于创新的混合策略：结构分层解析 + 语义增强翻译 + 格式智能重建
"""

import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import json
import re
from typing import Dict, List, Tuple, Any
import openai

class StructuralParser:
    """结构分层解析器 - 将文档分解为内容层、格式层、布局层"""
    
    def __init__(self):
        self.content_layer = []  # 纯文本内容
        self.format_layer = []   # 格式信息
        self.layout_layer = []   # 布局信息
        self.anchors = {}        # 锚点映射
    
    def parse_document(self, doc_path: str) -> Dict[str, Any]:
        """解析Word文档，提取三层信息"""
        try:
            doc = Document(doc_path)
            
            # 初始化解析结果
            result = {
                'content_layer': [],
                'format_layer': [],
                'layout_layer': [],
                'anchors': {},
                'metadata': {
                    'total_paragraphs': 0,
                    'total_tables': 0,
                    'total_images': 0,
                    'total_pages': 0
                }
            }
            
            # 解析段落
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # 内容层：纯文本
                    content_info = {
                        'id': f'para_{i}',
                        'text': paragraph.text.strip(),
                        'type': 'paragraph'
                    }
                    result['content_layer'].append(content_info)
                    
                    # 格式层：样式信息
                    format_info = {
                        'id': f'para_{i}',
                        'style': paragraph.style.name,
                        'alignment': str(paragraph.alignment),
                        'runs': []
                    }
                    
                    for run in paragraph.runs:
                        run_info = {
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': run.font.name,
                            'font_size': run.font.size.pt if run.font.size else None
                        }
                        format_info['runs'].append(run_info)
                    
                    result['format_layer'].append(format_info)
                    
                    # 布局层：结构信息
                    layout_info = {
                        'id': f'para_{i}',
                        'is_heading': paragraph.style.name.startswith('Heading'),
                        'heading_level': self._get_heading_level(paragraph.style.name),
                        'page_break_before': paragraph._element.getparent().tag.endswith('pPr')
                    }
                    result['layout_layer'].append(layout_info)
                    
                    result['metadata']['total_paragraphs'] += 1
            
            # 解析表格
            for i, table in enumerate(doc.tables):
                table_content = self._parse_table(table, i)
                result['content_layer'].extend(table_content['content'])
                result['format_layer'].extend(table_content['format'])
                result['layout_layer'].extend(table_content['layout'])
                result['metadata']['total_tables'] += 1
            
            # 解析图片
            for i, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    image_info = {
                        'id': f'img_{i}',
                        'type': 'image',
                        'target': rel.target_ref,
                        'content_type': rel.target_content_type
                    }
                    result['content_layer'].append(image_info)
                    result['metadata']['total_images'] += 1
            
            return result
            
        except Exception as e:
            st.error(f"文档解析失败: {str(e)}")
            return None
    
    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别"""
        if 'Heading 1' in style_name:
            return 1
        elif 'Heading 2' in style_name:
            return 2
        elif 'Heading 3' in style_name:
            return 3
        elif 'Heading 4' in style_name:
            return 4
        elif 'Heading 5' in style_name:
            return 5
        elif 'Heading 6' in style_name:
            return 6
        return 0
    
    def _parse_table(self, table, table_index: int) -> Dict[str, List]:
        """解析表格 - 修复重复问题"""
        content = []
        format_info = []
        layout_info = []
        
        # 使用集合来跟踪已处理的单元格内容，避免重复
        processed_cells = set()
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 跳过空单元格
                if not cell_text:
                    continue
                
                # 创建唯一标识符，避免重复处理相同内容
                cell_key = f"{table_index}_{row_idx}_{col_idx}_{cell_text}"
                if cell_key in processed_cells:
                    continue
                
                processed_cells.add(cell_key)
                
                cell_id = f'table_{table_index}_row_{row_idx}_col_{col_idx}'
                
                # 内容层
                content.append({
                    'id': cell_id,
                    'text': cell_text,
                    'type': 'table_cell',
                    'table_index': table_index,
                    'row': row_idx,
                    'col': col_idx
                })
                
                # 格式层
                format_info.append({
                    'id': cell_id,
                    'cell_style': 'table_cell',
                    'runs': []
                })
                
                # 布局层
                layout_info.append({
                    'id': cell_id,
                    'type': 'table_cell',
                    'table_index': table_index,
                    'row': row_idx,
                    'col': col_idx
                })
        
        return {
            'content': content,
            'format': format_info,
            'layout': layout_info
        }

class SemanticTranslator:
    """语义增强翻译器 - 支持上下文记忆、术语锁定、风格模仿、专有名词保护"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.context_memory = {}  # 上下文记忆
        self.terminology = {}     # 术语锁定
        self.style_examples = {}  # 风格示例
        self.proper_nouns = set()  # 专有名词集合
        self._init_proper_nouns()  # 初始化常见专有名词
        
    def set_terminology(self, terms: Dict[str, str]):
        """设置术语锁定"""
        self.terminology = terms
    
    def set_style_examples(self, examples: Dict[str, str]):
        """设置风格示例"""
        self.style_examples = examples
    
    def _init_proper_nouns(self):
        """初始化常见专有名词"""
        # 技术公司/平台
        tech_companies = [
            'GitHub', 'Google', 'Microsoft', 'Apple', 'Amazon', 'Facebook', 'Meta',
            'Twitter', 'LinkedIn', 'Instagram', 'YouTube', 'Netflix', 'Spotify',
            'OpenAI', 'Anthropic', 'Claude', 'ChatGPT', 'GPT', 'DALL-E',
            'Streamlit', 'Docker', 'Kubernetes', 'React', 'Vue', 'Angular',
            'Node.js', 'Python', 'JavaScript', 'TypeScript', 'Java', 'C++',
            'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy'
        ]
        
        # 开源项目
        open_source = [
            'Linux', 'Apache', 'Nginx', 'MySQL', 'PostgreSQL', 'MongoDB',
            'Redis', 'Elasticsearch', 'Kibana', 'Grafana', 'Prometheus',
            'Jenkins', 'GitLab', 'Bitbucket', 'Jira', 'Confluence'
        ]
        
        # 协议和标准
        protocols = [
            'HTTP', 'HTTPS', 'FTP', 'SSH', 'SMTP', 'POP3', 'IMAP',
            'TCP', 'UDP', 'IP', 'DNS', 'SSL', 'TLS', 'OAuth', 'JWT',
            'REST', 'GraphQL', 'WebSocket', 'gRPC', 'JSON', 'XML', 'YAML'
        ]
        
        # 学术机构
        universities = [
            'MIT', 'Stanford', 'Harvard', 'Berkeley', 'CMU', 'Oxford',
            'Cambridge', 'Yale', 'Princeton', 'Caltech', 'UCLA', 'NYU'
        ]
        
        # 添加到专有名词集合
        all_proper_nouns = tech_companies + open_source + protocols + universities
        self.proper_nouns.update(all_proper_nouns)
    
    def add_proper_nouns(self, nouns: List[str]):
        """添加自定义专有名词"""
        self.proper_nouns.update(nouns)
    
    def _protect_proper_nouns(self, text: str) -> Tuple[str, Dict[str, str]]:
        """保护专有名词，返回替换后的文本和映射表"""
        protected_text = text
        noun_mapping = {}
        
        # 按长度排序，优先匹配长专有名词
        sorted_nouns = sorted(self.proper_nouns, key=len, reverse=True)
        
        for noun in sorted_nouns:
            if noun in protected_text:
                # 创建占位符
                placeholder = f"__PROPER_NOUN_{len(noun_mapping)}__"
                noun_mapping[placeholder] = noun
                protected_text = protected_text.replace(noun, placeholder)
        
        return protected_text, noun_mapping
    
    def _identify_special_names_with_ai(self, text: str) -> List[str]:
        """使用OpenAI智能识别特殊名称（GitHub库名、项目名等）"""
        try:
            # 构建识别提示
            prompt = f"""
请仔细分析以下文本，识别出所有应该保持不变的特殊名称，包括但不限于：

1. GitHub库名（如：naiveHobo/InvoiceNet, microsoft/TypeScript）
2. 项目名称（如：React, Vue.js, Angular）
3. 技术框架名（如：TensorFlow, PyTorch, Scikit-learn）
4. 公司/组织名（如：Google, Microsoft, OpenAI）
5. 产品名称（如：ChatGPT, GitHub Copilot）
6. 协议/标准（如：HTTP, JSON, XML）
7. 编程语言（如：Python, JavaScript, TypeScript）

文本内容：{text}

请只返回识别出的特殊名称，每行一个，不要包含任何解释或其他内容。
如果文本中没有特殊名称，请返回空行。
"""
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "你是一个专业的文本分析助手，专门识别技术文档中的特殊名称。"},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.1
            )
            
            # 解析返回的特殊名称
            identified_names = []
            response_text = response.choices[0].message.content.strip()
            
            if response_text:
                for line in response_text.split('\n'):
                    name = line.strip()
                    if name and name in text:
                        identified_names.append(name)
            
            return identified_names
            
        except Exception as e:
            print(f"AI识别特殊名称失败: {str(e)}")
            return []
    
    def _protect_special_names_with_ai(self, text: str) -> Tuple[str, Dict[str, str]]:
        """使用AI智能保护特殊名称"""
        try:
            # 使用AI识别特殊名称
            special_names = self._identify_special_names_with_ai(text)
            
            protected_text = text
            noun_mapping = {}
            
            # 保护AI识别的特殊名称
            for name in special_names:
                if name in protected_text:
                    placeholder = f"__SPECIAL_NAME_{len(noun_mapping)}__"
                    noun_mapping[placeholder] = name
                    protected_text = protected_text.replace(name, placeholder)
            
            return protected_text, noun_mapping
            
        except Exception as e:
            print(f"AI保护特殊名称失败: {str(e)}")
            return text, {}
    
    def _restore_proper_nouns(self, text: str, noun_mapping: Dict[str, str]) -> str:
        """恢复专有名词"""
        restored_text = text
        for placeholder, noun in noun_mapping.items():
            restored_text = restored_text.replace(placeholder, noun)
        return restored_text
    
    def translate_with_context(self, content_items: List[Dict], target_lang: str) -> List[Dict]:
        """带上下文的翻译 - 修复重复内容问题"""
        try:
            openai.api_key = self.api_key
            
            # 构建上下文记忆
            context_prompt = self._build_context_prompt(content_items, target_lang)
            
            # 批量翻译，避免重复
            translated_items = []
            processed_texts = {}  # 用于避免重复翻译相同内容
            
            for item in content_items:
                if item['type'] == 'paragraph':
                    # 检查是否已经翻译过相同内容
                    text_key = item['text'].strip()
                    if text_key in processed_texts:
                        translated_text = processed_texts[text_key]
                    else:
                        translated_text = self._translate_paragraph(item, context_prompt, target_lang)
                        processed_texts[text_key] = translated_text
                    
                    translated_items.append({
                        **item,
                        'translated_text': translated_text
                    })
                    
                elif item['type'] == 'table_cell':
                    # 表格单元格特殊处理，避免重复翻译
                    text_key = f"table_{item.get('table_index', 0)}_{item.get('row', 0)}_{item.get('col', 0)}_{item['text'].strip()}"
                    if text_key in processed_texts:
                        translated_text = processed_texts[text_key]
                    else:
                        translated_text = self._translate_table_cell(item, context_prompt, target_lang)
                        processed_texts[text_key] = translated_text
                    
                    translated_items.append({
                        **item,
                        'translated_text': translated_text
                    })
                else:
                    translated_items.append(item)
            
            return translated_items
            
        except Exception as e:
            st.error(f"语义翻译失败: {str(e)}")
            return content_items
    
    def _build_context_prompt(self, content_items: List[Dict], target_lang: str) -> str:
        """构建上下文提示"""
        # 收集文档上下文
        context_texts = []
        for item in content_items[:10]:  # 取前10个段落作为上下文
            if item['type'] in ['paragraph', 'table_cell']:
                context_texts.append(item['text'])
        
        context = " ".join(context_texts)
        
        # 构建术语提示
        term_prompt = ""
        if self.terminology:
            term_prompt = f"\n术语锁定：{json.dumps(self.terminology, ensure_ascii=False)}"
        
        # 构建风格提示
        style_prompt = ""
        if self.style_examples:
            style_prompt = f"\n风格示例：{json.dumps(self.style_examples, ensure_ascii=False)}"
        
        return f"""
        上下文文档：{context}
        {term_prompt}
        {style_prompt}
        
        请将以下文本翻译为{target_lang}，保持专业术语一致性和文档风格。
        """
    
    def _translate_paragraph(self, item: Dict, context: str, target_lang: str) -> str:
        """翻译段落 - 带AI智能专有名词保护"""
        try:
            original_text = item['text']
            
            # 使用AI智能识别和保护特殊名称
            protected_text, ai_noun_mapping = self._protect_special_names_with_ai(original_text)
            
            # 同时使用传统专有名词保护
            protected_text, traditional_noun_mapping = self._protect_proper_nouns(protected_text)
            
            # 合并映射表
            combined_mapping = {**ai_noun_mapping, **traditional_noun_mapping}
            
            # 构建翻译提示，强调不要翻译特殊名称
            special_name_instruction = ""
            if combined_mapping:
                protected_names = list(combined_mapping.values())
                special_name_instruction = f"\n重要：请保持以下特殊名称不变：{', '.join(protected_names)}"
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"You are a professional document translator. {context}{special_name_instruction}"},
                    {"role": "user", "content": f"Translate this paragraph to {target_lang}: {protected_text}"}
                ],
                max_tokens=1000,
                temperature=0.1
            )
            
            translated_text = response.choices[0].message.content
            
            # 恢复所有特殊名称
            final_text = self._restore_proper_nouns(translated_text, combined_mapping)
            
            return final_text
        except Exception as e:
            st.warning(f"段落翻译失败: {str(e)}")
            return item['text']
    
    def _translate_table_cell(self, item: Dict, context: str, target_lang: str) -> str:
        """翻译表格单元格 - 带AI智能专有名词保护"""
        try:
            original_text = item['text']
            
            # 使用AI智能识别和保护特殊名称
            protected_text, ai_noun_mapping = self._protect_special_names_with_ai(original_text)
            
            # 同时使用传统专有名词保护
            protected_text, traditional_noun_mapping = self._protect_proper_nouns(protected_text)
            
            # 合并映射表
            combined_mapping = {**ai_noun_mapping, **traditional_noun_mapping}
            
            # 构建翻译提示，强调不要翻译特殊名称
            special_name_instruction = ""
            if combined_mapping:
                protected_names = list(combined_mapping.values())
                special_name_instruction = f"\n重要：请保持以下特殊名称不变：{', '.join(protected_names)}"
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"You are a professional document translator. {context}{special_name_instruction}"},
                    {"role": "user", "content": f"Translate this table cell content to {target_lang}: {protected_text}"}
                ],
                max_tokens=500,
                temperature=0.1
            )
            
            translated_text = response.choices[0].message.content
            
            # 恢复所有特殊名称
            final_text = self._restore_proper_nouns(translated_text, combined_mapping)
            
            return final_text
        except Exception as e:
            st.warning(f"表格单元格翻译失败: {str(e)}")
            return item['text']

class SmartReconstructor:
    """格式智能重建器 - 利用锚点映射重组文档"""
    
    def __init__(self):
        self.anchors = {}
        self.format_preservation = True
    
    def reconstruct_document(self, original_doc_path: str, translated_content: List[Dict], 
                           format_layer: List[Dict], layout_layer: List[Dict], 
                           output_path: str) -> bool:
        """重建文档"""
        try:
            # 加载原文档
            doc = Document(original_doc_path)
            
            # 创建翻译映射
            translation_map = {item['id']: item['translated_text'] for item in translated_content 
                              if 'translated_text' in item}
            
            # 重建段落
            self._reconstruct_paragraphs(doc, translation_map, format_layer)
            
            # 重建表格
            self._reconstruct_tables(doc, translation_map, format_layer)
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            st.error(f"文档重建失败: {str(e)}")
            return False
    
    def _reconstruct_paragraphs(self, doc: Document, translation_map: Dict, format_layer: List[Dict]):
        """重建段落"""
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_id = f'para_{i}'
                if para_id in translation_map:
                    # 智能处理翻译长度变化
                    original_text = paragraph.text
                    translated_text = translation_map[para_id]
                    
                    # 如果翻译长度变化不大，直接替换
                    if abs(len(translated_text) - len(original_text)) / len(original_text) < 0.5:
                        paragraph.text = translated_text
                    else:
                        # 长度变化大，需要智能调整
                        self._smart_text_replacement(paragraph, translated_text)
    
    def _reconstruct_tables(self, doc: Document, translation_map: Dict, format_layer: List[Dict]):
        """重建表格 - 修复重复问题"""
        table_index = 0
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # 创建单元格ID
                    cell_id = f'table_{table_index}_row_{row_idx}_col_{col_idx}'
                    
                    # 查找对应的翻译
                    if cell_id in translation_map:
                        translated_text = translation_map[cell_id]
                        if translated_text and translated_text.strip():
                            # 智能处理翻译长度变化
                            original_text = cell.text.strip()
                            if original_text:
                                # 如果翻译长度变化不大，直接替换
                                if abs(len(translated_text) - len(original_text)) / len(original_text) < 0.5:
                                    cell.text = translated_text
                                else:
                                    # 长度变化大，需要智能调整
                                    self._smart_cell_replacement(cell, translated_text)
            table_index += 1
    
    def _smart_cell_replacement(self, cell, translated_text: str):
        """智能单元格文本替换"""
        # 清空单元格内容
        cell.text = ""
        # 添加翻译文本
        cell.text = translated_text
    
    def _smart_text_replacement(self, paragraph, translated_text: str):
        """智能文本替换，处理长度变化"""
        # 保持原有的run结构，但调整内容
        runs = paragraph.runs
        if runs:
            # 将翻译文本分配给第一个run
            runs[0].text = translated_text
            # 清空其他run
            for run in runs[1:]:
                run.text = ""

class FormatCorrector:
    """格式纠错模块 - 检测和修复排版问题"""
    
    def __init__(self):
        self.correction_rules = []
    
    def detect_format_issues(self, doc_path: str) -> List[Dict]:
        """检测格式问题"""
        issues = []
        try:
            doc = Document(doc_path)
            
            # 检测表格溢出
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if len(cell.text) > 100:  # 单元格内容过长
                            issues.append({
                                'type': 'table_overflow',
                                'location': 'table',
                                'description': '表格单元格内容过长',
                                'suggestion': '建议拆分长文本'
                            })
            
            # 检测标题格式
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    if not paragraph.text.strip():
                        issues.append({
                            'type': 'empty_heading',
                            'location': 'paragraph',
                            'description': '空标题',
                            'suggestion': '建议添加标题内容或删除空标题'
                        })
            
            return issues
            
        except Exception as e:
            st.error(f"格式检测失败: {str(e)}")
            return []
    
    def auto_fix_issues(self, doc_path: str, issues: List[Dict]) -> bool:
        """自动修复格式问题"""
        try:
            doc = Document(doc_path)
            
            for issue in issues:
                if issue['type'] == 'table_overflow':
                    # 自动拆分长文本
                    self._fix_table_overflow(doc)
                elif issue['type'] == 'empty_heading':
                    # 删除空标题
                    self._fix_empty_headings(doc)
            
            doc.save(doc_path)
            return True
            
        except Exception as e:
            st.error(f"自动修复失败: {str(e)}")
            return False
    
    def _fix_table_overflow(self, doc: Document):
        """修复表格溢出"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.text) > 100:
                        # 拆分长文本
                        text = cell.text
                        if len(text) > 100:
                            # 在合适的位置拆分
                            split_point = text.rfind(' ', 0, 100)
                            if split_point > 0:
                                cell.text = text[:split_point] + '\n' + text[split_point+1:]
    
    def _fix_empty_headings(self, doc: Document):
        """修复空标题"""
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and not paragraph.text.strip():
                paragraphs_to_remove.append(paragraph)
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

class DualViewEditor:
    """双视图编辑器 - 左右对比显示"""
    
    def __init__(self):
        self.original_content = []
        self.translated_content = []
    
    def display_dual_view(self, original_items: List[Dict], translated_items: List[Dict]):
        """显示双视图 - 修复重复内容问题"""
        st.subheader("📖 双视图编辑器")
        
        # 去重处理
        original_unique = self._deduplicate_items(original_items)
        translated_unique = self._deduplicate_items(translated_items)
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 📄 原文")
            self._display_content(original_unique, "original")
        
        with col2:
            st.markdown("### 🌐 译文")
            self._display_content(translated_unique, "translated")
    
    def _deduplicate_items(self, items: List[Dict]) -> List[Dict]:
        """去重处理，避免重复显示"""
        seen_texts = set()
        unique_items = []
        
        for item in items:
            text_key = item.get('text', '').strip()
            if text_key and text_key not in seen_texts:
                seen_texts.add(text_key)
                unique_items.append(item)
            elif not text_key:  # 保留空内容项
                unique_items.append(item)
        
        return unique_items
    
    def _display_content(self, items: List[Dict], view_type: str):
        """显示内容 - 优化显示逻辑"""
        # 按类型分组显示
        paragraphs = [item for item in items if item['type'] == 'paragraph']
        table_cells = [item for item in items if item['type'] == 'table_cell']
        
        # 显示段落
        if paragraphs:
            st.markdown("**📝 段落内容:**")
            for i, item in enumerate(paragraphs):
                if item['text'].strip():  # 只显示非空段落
                    if view_type == "translated" and 'translated_text' in item:
                        st.text_area(f"段落 {i+1}", value=item['translated_text'], height=80, key=f"para_{view_type}_{i}")
                    else:
                        st.text_area(f"段落 {i+1}", value=item['text'], height=80, key=f"para_{view_type}_{i}")
        
        # 显示表格单元格 - 修复重复问题
        if table_cells:
            st.markdown("**📊 表格内容:**")
            # 按表格分组
            tables = {}
            for item in table_cells:
                table_idx = item.get('table_index', 0)
                if table_idx not in tables:
                    tables[table_idx] = []
                tables[table_idx].append(item)
            
            for table_idx, cells in tables.items():
                st.markdown(f"**表格 {table_idx + 1}:**")
                
                # 去重处理：使用集合跟踪已显示的单元格
                displayed_cells = set()
                
                # 按行列组织
                rows = {}
                for cell in cells:
                    row = cell.get('row', 0)
                    col = cell.get('col', 0)
                    cell_text = cell['text'].strip()
                    
                    # 跳过空单元格和重复内容
                    if not cell_text or cell_text in displayed_cells:
                        continue
                    
                    displayed_cells.add(cell_text)
                    
                    if row not in rows:
                        rows[row] = {}
                    rows[row][col] = cell
                
                # 显示表格内容
                for row_idx in sorted(rows.keys()):
                    cols = rows[row_idx]
                    if cols:  # 只显示有内容的行
                        st.markdown(f"**第 {row_idx + 1} 行:**")
                        for col_idx in sorted(cols.keys()):
                            cell = cols[col_idx]
                            if cell['text'].strip():  # 只显示非空单元格
                                if view_type == "translated" and 'translated_text' in cell:
                                    st.text_input(f"列{col_idx+1}", value=cell['translated_text'], key=f"cell_{view_type}_{table_idx}_{row_idx}_{col_idx}")
                                else:
                                    st.text_input(f"列{col_idx+1}", value=cell['text'], key=f"cell_{view_type}_{table_idx}_{row_idx}_{col_idx}")

# 主应用类
class SmartDocumentTranslator:
    """智能文档翻译与格式保真系统主类"""
    
    def __init__(self):
        self.parser = StructuralParser()
        self.translator = None
        self.reconstructor = SmartReconstructor()
        self.corrector = FormatCorrector()
        self.editor = DualViewEditor()
    
    def set_translator(self, api_key: str):
        """设置翻译器"""
        self.translator = SemanticTranslator(api_key)
    
    def process_document(self, doc_path: str, target_lang: str, output_path: str) -> bool:
        """处理文档的完整流程"""
        try:
            # 1. 结构分层解析
            st.info("🔍 正在进行结构分层解析...")
            parsed_doc = self.parser.parse_document(doc_path)
            if not parsed_doc:
                return False
            
            # 2. 语义增强翻译
            st.info("🤖 正在进行语义增强翻译...")
            if not self.translator:
                st.error("请先设置翻译器")
                return False
            
            translated_content = self.translator.translate_with_context(
                parsed_doc['content_layer'], target_lang
            )
            
            # 3. 格式智能重建
            st.info("🔧 正在进行格式智能重建...")
            success = self.reconstructor.reconstruct_document(
                doc_path, translated_content, 
                parsed_doc['format_layer'], parsed_doc['layout_layer'], 
                output_path
            )
            
            if success:
                # 4. 格式纠错
                st.info("🔍 正在进行格式纠错...")
                issues = self.corrector.detect_format_issues(output_path)
                if issues:
                    st.warning(f"发现 {len(issues)} 个格式问题，正在自动修复...")
                    self.corrector.auto_fix_issues(output_path, issues)
                
                return True
            
            return False
            
        except Exception as e:
            st.error(f"文档处理失败: {str(e)}")
            return False
