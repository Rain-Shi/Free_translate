"""
智能文档查看器演示
"""

import streamlit as st
from document_viewer import DocumentViewer
from docx import Document
import tempfile
import os

def create_demo_documents():
    """创建演示文档"""
    # 创建原文文档
    original_doc = Document()
    original_doc.add_heading('Meeting Report', 0)
    original_doc.add_heading('Introduction', level=1)
    original_doc.add_paragraph('This is the first paragraph of the meeting report.')
    original_doc.add_paragraph('The meeting was held on September 28, 2025.')
    original_doc.add_paragraph('Key topics discussed include project planning and resource allocation.')
    
    original_doc.add_heading('Main Discussion', level=1)
    original_doc.add_paragraph('The team discussed various technical challenges.')
    original_doc.add_paragraph('GitHub integration was a major topic.')
    original_doc.add_paragraph('OpenAI API usage was also discussed.')
    
    original_doc.add_heading('Action Items', level=1)
    original_doc.add_paragraph('Complete the project documentation.')
    original_doc.add_paragraph('Review the code changes.')
    original_doc.add_paragraph('Schedule the next meeting.')
    
    # 创建译文文档
    translated_doc = Document()
    translated_doc.add_heading('会议报告', 0)
    translated_doc.add_heading('介绍', level=1)
    translated_doc.add_paragraph('这是会议报告的第一段。')
    translated_doc.add_paragraph('会议于2025年9月28日举行。')
    translated_doc.add_paragraph('讨论的主要话题包括项目规划和资源分配。')
    
    translated_doc.add_heading('主要讨论', level=1)
    translated_doc.add_paragraph('团队讨论了各种技术挑战。')
    translated_doc.add_paragraph('GitHub集成是一个主要话题。')
    translated_doc.add_paragraph('还讨论了OpenAI API的使用。')
    
    translated_doc.add_heading('行动项目', level=1)
    translated_doc.add_paragraph('完成项目文档。')
    translated_doc.add_paragraph('审查代码更改。')
    translated_doc.add_paragraph('安排下次会议。')
    
    # 保存到临时文件
    original_path = tempfile.mktemp(suffix='.docx')
    translated_path = tempfile.mktemp(suffix='.docx')
    
    original_doc.save(original_path)
    translated_doc.save(translated_path)
    
    return original_path, translated_path

def main():
    """主演示函数"""
    st.set_page_config(
        page_title="智能文档查看器演示",
        page_icon="📖",
        layout="wide"
    )
    
    st.title("📖 智能文档查看器演示")
    st.markdown("---")
    
    # 创建演示文档
    if st.button("🚀 创建演示文档"):
        with st.spinner("正在创建演示文档..."):
            original_path, translated_path = create_demo_documents()
            
            # 初始化文档查看器
            viewer = DocumentViewer()
            
            # 加载文档
            if viewer.load_documents(original_path, translated_path):
                st.success("✅ 演示文档创建成功！")
                
                # 显示文档摘要
                viewer.display_document_summary()
                
                # 显示文档查看器
                viewer.display_document_viewer()
                
                # 清理临时文件
                try:
                    os.unlink(original_path)
                    os.unlink(translated_path)
                except:
                    pass
            else:
                st.error("❌ 演示文档创建失败")
    
    # 功能说明
    st.markdown("---")
    st.subheader("🎯 功能特性")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 📄 页面导航")
        st.markdown("- 统一翻页控制")
        st.markdown("- 页面跳转功能")
        st.markdown("- 上一页/下一页")
    
    with col2:
        st.markdown("### 🔍 段落对比")
        st.markdown("- 点击段落查看对比")
        st.markdown("- 详细统计信息")
        st.markdown("- 编辑功能")
    
    with col3:
        st.markdown("### 📊 文档摘要")
        st.markdown("- 总段落数统计")
        st.markdown("- 页数统计")
        st.markdown("- 平均字数")

if __name__ == "__main__":
    main()
