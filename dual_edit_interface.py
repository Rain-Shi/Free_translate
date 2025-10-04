"""
左右分开的编辑界面
"""

import streamlit as st
from docx import Document
from typing import List, Dict, Any
import tempfile
import os

class DualEditInterface:
    """左右分开的编辑界面"""
    
    def __init__(self):
        self.original_paragraphs = []
        self.translated_paragraphs = []
        self.edited_paragraphs = []
    
    def load_documents(self, original_path: str, translated_path: str):
        """加载原文档和翻译文档"""
        try:
            # 读取原文档
            original_doc = Document(original_path)
            self.original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
            
            # 读取翻译文档
            translated_doc = Document(translated_path)
            self.translated_paragraphs = [p.text.strip() for p in translated_doc.paragraphs if p.text.strip()]
            
            # 初始化编辑后的段落
            self.edited_paragraphs = self.translated_paragraphs.copy()
            
            st.success("✅ 文档加载成功！")
            return True
        except Exception as e:
            st.error(f"❌ 文档加载失败: {str(e)}")
            return False
    
    def display_dual_edit_interface(self):
        """显示左右分开的编辑界面"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            st.warning("⚠️ 请先加载文档")
            return
        
        st.markdown("---")
        st.subheader("📝 左右编辑界面")
        
        # 显示统计信息
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("原文档段落数", len(self.original_paragraphs))
        
        with col2:
            st.metric("译文段落数", len(self.translated_paragraphs))
        
        with col3:
            st.metric("可编辑段落数", len(self.edited_paragraphs))
        
        # 一次性显示所有段落的左右编辑界面
        max_paragraphs = min(len(self.original_paragraphs), len(self.translated_paragraphs))
        
        if max_paragraphs > 0:
            # 显示所有段落的左右编辑界面
            for i in range(max_paragraphs):
                self._display_paragraph_edit_interface(i)
    
    def _display_paragraph_edit_interface(self, para_index: int):
        """显示单个段落的左右编辑界面"""
        st.markdown("---")
        st.markdown(f"### ✏️ 段落 {para_index + 1}")
        
        # 创建两列布局
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📝 原文**")
            original_text = self.original_paragraphs[para_index] if para_index < len(self.original_paragraphs) else ""
            st.text_area(
                "原文内容",
                value=original_text,
                height=150,
                key=f"original_display_{para_index}",
                disabled=True
            )
            
            # 原文统计
            st.markdown(f"**字数**: {len(original_text)}")
        
        with col2:
            st.markdown("**🌐 译文 (可编辑)**")
            # 使用session_state来保持编辑状态
            edit_key = f"edited_text_{para_index}"
            if edit_key not in st.session_state:
                st.session_state[edit_key] = self.edited_paragraphs[para_index] if para_index < len(self.edited_paragraphs) else ""
            
            edited_text = st.text_area(
                "译文内容 (可编辑)",
                value=st.session_state[edit_key],
                height=150,
                key=f"translated_edit_{para_index}",
                help="您可以在这里编辑译文内容"
            )
            
            # 更新编辑后的段落
            if edited_text != st.session_state[edit_key]:
                st.session_state[edit_key] = edited_text
                self.edited_paragraphs[para_index] = edited_text
            
            # 译文统计
            st.markdown(f"**字数**: {len(edited_text)}")
        
        # 对比统计
        if original_text and edited_text:
            self._display_edit_comparison_stats(original_text, edited_text)
        
        # 编辑操作按钮
        self._display_edit_actions(para_index)
    
    def _display_edit_comparison_stats(self, original_text: str, edited_text: str):
        """显示编辑对比统计"""
        st.markdown("---")
        st.markdown("### 📊 编辑对比统计")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            length_ratio = len(edited_text) / len(original_text) if original_text else 1
            st.metric("长度比例", f"{length_ratio:.2f}")
        
        with col2:
            word_count_orig = len(original_text.split())
            word_count_edit = len(edited_text.split())
            st.metric("原文词数", word_count_orig)
        
        with col3:
            st.metric("编辑后词数", word_count_edit)
        
        with col4:
            if word_count_orig > 0:
                word_ratio = word_count_edit / word_count_orig
                st.metric("词数比例", f"{word_ratio:.2f}")
    
    def _display_edit_actions(self, para_index: int):
        """显示编辑操作按钮"""
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 重置为原文", key=f"reset_{para_index}"):
                original_text = self.original_paragraphs[para_index] if para_index < len(self.original_paragraphs) else ""
                st.session_state[f"edited_text_{para_index}"] = original_text
                self.edited_paragraphs[para_index] = original_text
                st.success("✅ 已重置为原文")
                st.rerun()
        
        with col2:
            if st.button("🔄 重置为译文", key=f"reset_trans_{para_index}"):
                translated_text = self.translated_paragraphs[para_index] if para_index < len(self.translated_paragraphs) else ""
                st.session_state[f"edited_text_{para_index}"] = translated_text
                self.edited_paragraphs[para_index] = translated_text
                st.success("✅ 已重置为译文")
                st.rerun()
        
        with col3:
            if st.button("📋 复制原文", key=f"copy_{para_index}"):
                st.session_state[f"edited_text_{para_index}"] = self.original_paragraphs[para_index]
                self.edited_paragraphs[para_index] = self.original_paragraphs[para_index]
                st.success("✅ 已复制原文")
                st.rerun()
    
    def display_all_paragraphs_edit(self):
        """显示所有段落的编辑界面"""
        if not self.original_paragraphs or not self.translated_paragraphs:
            return
        
        st.markdown("---")
        st.subheader("📋 所有段落编辑")
        
        # 创建可滚动的段落编辑列表
        max_paragraphs = min(len(self.original_paragraphs), len(self.translated_paragraphs))
        
        for i in range(max_paragraphs):
            with st.expander(f"段落 {i+1}", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**原文**")
                    st.text_area(
                        f"原文_{i}",
                        value=self.original_paragraphs[i],
                        height=100,
                        key=f"original_all_{i}",
                        disabled=True
                    )
                
                with col2:
                    st.markdown("**译文 (可编辑)**")
                    edit_key = f"edited_all_{i}"
                    if edit_key not in st.session_state:
                        st.session_state[edit_key] = self.edited_paragraphs[i] if i < len(self.edited_paragraphs) else ""
                    
                    edited_text = st.text_area(
                        f"译文_{i}",
                        value=st.session_state[edit_key],
                        height=100,
                        key=f"translated_all_{i}",
                        help="您可以在这里编辑译文内容"
                    )
                    
                    # 更新编辑后的段落
                    if edited_text != st.session_state[edit_key]:
                        st.session_state[edit_key] = edited_text
                        self.edited_paragraphs[i] = edited_text
    
    def create_final_document(self, output_path: str):
        """创建最终文档"""
        try:
            # 读取原文档结构
            original_doc = Document()
            
            # 创建新文档
            new_doc = Document()
            
            # 添加标题
            new_doc.add_heading('翻译后的文档', 0)
            
            # 添加编辑后的段落
            for i, edited_text in enumerate(self.edited_paragraphs):
                if edited_text.strip():
                    new_doc.add_paragraph(edited_text)
            
            # 保存文档
            new_doc.save(output_path)
            return True
        except Exception as e:
            st.error(f"❌ 创建最终文档失败: {str(e)}")
            return False
    
    def get_edit_summary(self):
        """获取编辑摘要"""
        if not self.original_paragraphs or not self.edited_paragraphs:
            return {}
        
        # 计算修改统计
        modified_count = 0
        for i in range(min(len(self.translated_paragraphs), len(self.edited_paragraphs))):
            if self.translated_paragraphs[i] != self.edited_paragraphs[i]:
                modified_count += 1
        
        return {
            'total_paragraphs': len(self.edited_paragraphs),
            'modified_paragraphs': modified_count,
            'modification_rate': modified_count / len(self.edited_paragraphs) if self.edited_paragraphs else 0
        }
    
    def display_edit_summary(self):
        """显示编辑摘要"""
        summary = self.get_edit_summary()
        
        if summary:
            st.markdown("### 📊 编辑摘要")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("总段落数", summary['total_paragraphs'])
            
            with col2:
                st.metric("已修改段落数", summary['modified_paragraphs'])
            
            with col3:
                st.metric("修改率", f"{summary['modification_rate']:.1%}")
